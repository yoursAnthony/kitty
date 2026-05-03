"""Генерация курсовой работы по проекту Kittygram (docx).

Структура — по «Методическим рекомендациям по курсовой работе»:
    Введение → 5 глав → Заключение → Источники → Приложения.

Скриншоты: docs/screenshots/creative/ (12 PNG из творческого).
Диаграммы: docs/diagrams/{usecase_base, usecase_chat, deployment}.png.

Запуск:
    .venv/Scripts/python scripts/build_kursovaya_report.py
Результат: ../../Any/PREU/.../Курсовая_работа.docx
Затем применить титульник:
    .venv/Scripts/python scripts/insert_title_page.py kursovaya
"""
import io
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPO = Path(__file__).resolve().parent.parent
SCREENS = REPO / 'docs' / 'screenshots' / 'creative'
DIAGRAMS = REPO / 'docs' / 'diagrams'
OUT = (
    REPO.parent.parent.parent
    / 'Any' / 'PREU'
    / 'интеграция и управление приложениями на удаленном сервере'
    / 'Курсовая_работа.docx'
)

ACCESS_DATE = '03.05.2026'


# ─────────── Хелперы оформления ───────────

def add_para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             indent=True, size=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    pf = p.paragraph_format
    pf.first_line_indent = Cm(1.25) if indent else Cm(0)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p.alignment = align
    return p


def add_heading_centered(doc, text, page_break_before=False):
    p = doc.add_paragraph()
    if page_break_before:
        run0 = p.add_run()
        run0.add_break(WD_BREAK.PAGE)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    return p


def add_heading_chapter(doc, text):
    """Глава курсовой — с разрыва страницы, по центру, 14pt bold UPPERCASE."""
    return add_heading_centered(doc, text, page_break_before=True)


def add_heading_section(doc, text):
    """Раздел внутри главы — слева, 14pt bold."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.5
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)
    return p


def add_heading_subsection(doc, text):
    """Подраздел — слева, 14pt bold italic."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(2)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0.5)
    pf.line_spacing = 1.15
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    return p


def style_table_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=12):
    cell.text = ''
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    p.alignment = align
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.15
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        style_table_cell(table.rows[0].cells[j], h, bold=True,
                         align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            style_table_cell(table.rows[i].cells[j], str(val))
    if col_widths_cm:
        for row in table.rows:
            for j, w in enumerate(col_widths_cm):
                row.cells[j].width = Cm(w)
    return table


def add_table_caption(doc, num, text):
    """Подпись таблицы — над таблицей, слева, курсив."""
    p = doc.add_paragraph()
    run = p.add_run(f'Таблица {num} — {text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.15
    pf.space_before = Pt(6)
    pf.space_after = Pt(2)
    return p


def add_image(doc, path, width_cm=14.0, caption=None):
    if not path.exists():
        add_para(doc, f'[ОТСУТСТВУЕТ ФАЙЛ СКРИНА: {path.name}]', bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.first_line_indent = Cm(0)
        c.paragraph_format.line_spacing = 1.15
        c.paragraph_format.space_after = Pt(6)
        cr = c.add_run(caption)
        cr.italic = True
        cr.font.name = 'Times New Roman'
        cr.font.size = Pt(12)


def add_hyperlink(paragraph, url, text):
    """Кликабельная гиперссылка в существующем параграфе (TNR 14, синий, подчёркнутый)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rPr.append(rFonts)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '28')  # 14pt = 28 half-points
    rPr.append(sz)

    new_run.append(rPr)

    text_el = OxmlElement('w:t')
    text_el.text = text
    text_el.set(qn('xml:space'), 'preserve')
    new_run.append(text_el)

    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)
    return hyperlink


def add_source(doc, num, parts):
    """Источник списка литературы. parts — список: либо строка, либо tuple ('link', url, text).

    Пример:
        add_source(doc, 1, [
            'Django Documentation [Электронный ресурс]. — URL: ',
            ('link', 'https://docs.djangoproject.com/en/4.2/', 'https://docs.djangoproject.com/en/4.2/'),
            f' (дата обращения: {ACCESS_DATE}).',
        ])
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0.5)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    prefix = p.add_run(f'{num}. ')
    prefix.font.name = 'Times New Roman'
    prefix.font.size = Pt(14)

    for part in parts:
        if isinstance(part, tuple) and part[0] == 'link':
            add_hyperlink(p, part[1], part[2])
        else:
            r = p.add_run(part)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)
    return p


def add_bullet(doc, text):
    """Маркированный пункт через дефис (для совместимости со стилями РЭУ)."""
    p = doc.add_paragraph()
    run = p.add_run('— ' + text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(1.25)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_numbered(doc, num, text):
    """Нумерованный пункт (1) текст)."""
    p = doc.add_paragraph()
    run = p.add_run(f'{num}) ' + text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(1.25)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def setup_page(doc):
    section = doc.sections[0]
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')

    pPr = style.element.get_or_add_pPr()
    spc = pPr.find(qn('w:spacing'))
    if spc is None:
        spc = OxmlElement('w:spacing')
        pPr.append(spc)
    spc.set(qn('w:line'), '360')
    spc.set(qn('w:lineRule'), 'auto')


# ─────────── Сборка документа ───────────

def build():
    doc = Document()
    setup_page(doc)

    # ───── ШАПКА (черновая, заменяется через insert_title_page.py kursovaya) ─────
    add_heading_centered(doc, 'КУРСОВАЯ РАБОТА')
    add_para(doc, 'по дисциплине «Интеграция и управление приложениями '
                  'на удалённом сервере».', indent=False,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, 'Тема: «Проектирование и реализация серверной части проекта '
                  'Kittygram для поддержки пользовательского сценария '
                  '"Мини-чат по заявкам"».', indent=False,
             align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_para(doc, 'Студент: Сергеев А. А., группа ПИ2у/24б.', indent=False,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, 'Преподаватель: Брызгалов А. А.', indent=False,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, 'Москва, 2026', indent=False,
             align=WD_ALIGN_PARAGRAPH.CENTER)

    # ───── СОДЕРЖАНИЕ (статическое — TOC field в Word всё равно требует Update) ─────
    add_heading_chapter(doc, 'СОДЕРЖАНИЕ')
    toc = [
        'Введение ............................................................ 3',
        'Глава 1. Аналитика и постановка задачи ............................. 5',
        '   1.1. Описание предметной области ................................ 5',
        '   1.2. Функциональные требования .................................. 6',
        '   1.3. Нефункциональные требования ................................ 7',
        '   1.4. Роли пользователей ......................................... 8',
        '   1.5. Ограничения и допущения .................................... 9',
        'Глава 2. Проектирование ........................................... 10',
        '   2.1. Архитектура приложения .................................... 10',
        '   2.2. Use Case диаграммы ........................................ 11',
        '   2.3. Модели данных ............................................. 13',
        '   2.4. API-контракт .............................................. 15',
        '   2.5. Права доступа ............................................. 16',
        '   2.6. Бизнес-валидации .......................................... 17',
        'Глава 3. Реализация проекта ....................................... 18',
        '   3.1. Структура проекта ......................................... 18',
        '   3.2. Стек реализации ........................................... 19',
        '   3.3. Реализация моделей ........................................ 20',
        '   3.4. Реализация ViewSet‘ов ................................. 21',
        '   3.5. Сериализаторы ............................................. 22',
        '   3.6. Permissions ............................................... 23',
        '   3.7. Сигналы автосоздания диалога .............................. 23',
        '   3.8. Фильтрация, поиск, пагинация ............................... 24',
        '   3.9. Обработка ошибок .......................................... 24',
        '   3.10. Документация API ......................................... 25',
        'Глава 4. Тестирование и проверка работоспособности ................ 26',
        '   4.1. Подход к тестированию ..................................... 26',
        '   4.2. Postman-коллекция ......................................... 26',
        '   4.3. Положительные сценарии .................................... 27',
        '   4.4. Негативные сценарии ....................................... 28',
        '   4.5. Прогон через newman ....................................... 28',
        'Глава 5. Развёртывание проекта .................................... 29',
        '   5.1. Контейнеризация (Dockerfile) ............................... 29',
        '   5.2. Многоконтейнерная архитектура (docker-compose) ............ 29',
        '   5.3. Сетевое взаимодействие .................................... 30',
        '   5.4. Хранение данных (volumes) ................................. 30',
        '   5.5. Конфигурация (.env) ....................................... 31',
        '   5.6. Deployment диаграмма ...................................... 31',
        '   5.7. Команды развёртывания ..................................... 32',
        'Заключение ........................................................ 33',
        'Список использованных источников .................................. 34',
        'Приложение А — Postman-коллекция .................................. 36',
        'Приложение Б — Полный API-контракт ................................ 37',
        'Приложение В — Конфигурации Docker ................................ 38',
    ]
    for line in toc:
        add_para(doc, line, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)

    # ╔══════════════════════════════════════════╗
    # ║         ВВЕДЕНИЕ                         ║
    # ╚══════════════════════════════════════════╝
    add_heading_chapter(doc, 'ВВЕДЕНИЕ')

    add_para(
        doc,
        'Современная веб-разработка предъявляет к серверной части веб-приложений '
        'высокие требования по сопровождаемости, безопасности и удобству '
        'развёртывания. Архитектурный стиль REST, предложенный Р. Филдингом '
        'в 2000 году, давно утвердился в качестве де-факто стандарта для построения '
        'API между клиентом и сервером. На стороне Python наиболее зрелым и широко '
        'распространённым решением для реализации REST API является Django REST '
        'Framework, надстраиваемый поверх веб-фреймворка Django. Такой стек '
        'позволяет достичь высокой скорости разработки, не жертвуя качеством кода '
        'и безопасностью, и широко используется в индустрии.'
    )
    add_para(
        doc,
        'Учебный проект Kittygram, рассматриваемый в рамках дисциплины '
        '«Интеграция и управление приложениями на удалённом сервере», представляет '
        'собой минимальный каркас социального сервиса для владельцев котов. '
        'В исходном виде он содержит одну модель кота и небольшой набор '
        'функциональных представлений на Django. Цель курсовой работы — '
        'превратить этот каркас в полноценный API-сервис уровня младшего '
        'production-приложения и расширить его пользовательским сценарием '
        '«мини-чат по заявкам на участие в кото-ивентах», в котором '
        'организатор события и желающий присоединиться пользователь могут '
        'обсудить детали участия по REST-эндпоинту чата.'
    )

    add_heading_section(doc, 'Актуальность темы')
    add_para(
        doc,
        'Актуальность работы определяется тем, что современная разработка '
        'распределённых веб-приложений требует уверенного владения связкой '
        '«Python + Django + DRF + PostgreSQL + Docker», понимания принципов '
        'построения REST API, а также практических навыков развёртывания такой '
        'системы под управлением docker-compose. Все перечисленные технологии '
        'являются базовыми для современного бэкенд-разработчика и востребованы '
        'на рынке труда. Работа выстроена так, чтобы охватить полный цикл — '
        'от анализа предметной области и проектирования моделей данных до '
        'тестирования через Postman и развёртывания в Docker.'
    )

    add_heading_section(doc, 'Объект и предмет исследования')
    add_para(
        doc,
        'Объект исследования — серверная часть веб-приложения для каталогизации '
        'котов и взаимодействия их владельцев в рамках событий-ивентов.'
    )
    add_para(
        doc,
        'Предмет исследования — методы и средства проектирования и реализации '
        'REST API на Django REST Framework, обеспечивающего поддержку '
        'пользовательского сценария «мини-чат по заявкам на кото-ивенты», '
        'с применением JSON Web Token для аутентификации, документации API через '
        'OpenAPI/Swagger UI/ReDoc и развёртыванием в многоконтейнерной '
        'docker-compose-инсталляции.'
    )

    add_heading_section(doc, 'Цель и задачи работы')
    add_para(
        doc,
        'Цель работы — спроектировать и реализовать серверную часть проекта '
        'Kittygram с расширением «мини-чат по заявкам на кото-ивенты», '
        'удовлетворяющим современным требованиям к API-сервисам.'
    )
    add_para(doc, 'Для достижения поставленной цели сформулированы следующие задачи:',
             indent=False)
    tasks = [
        'проанализировать предметную область, выделить роли пользователей '
        'и сформулировать функциональные и нефункциональные требования;',
        'спроектировать модели данных базового Kittygram (коты, теги, достижения) '
        'и расширения «мини-чат по заявкам» (ивенты, заявки, диалоги, сообщения);',
        'построить Use Case диаграммы для базовой части и для расширения, '
        'задокументировать API-контракт, права доступа и бизнес-валидации;',
        'реализовать REST API на Django 4.2 LTS и DRF 3.14 с применением '
        'ViewSet’ов, кастомных action, сериализаторов с валидациями, '
        'фильтрации, поиска и пагинации;',
        'настроить аутентификацию по JSON Web Token (через библиотеки Djoser '
        'и djangorestframework-simplejwt) и документацию OpenAPI '
        '(через drf-spectacular);',
        'провести функциональное тестирование API через Postman и newman, '
        'покрыть позитивные и негативные сценарии;',
        'реализовать развёртывание в Docker через docker-compose '
        '(сервисы: PostgreSQL, backend, nginx) и подготовить Deployment-диаграмму.',
    ]
    for i, t in enumerate(tasks, start=1):
        add_numbered(doc, i, t)

    add_heading_section(doc, 'Применяемые технологии и инструменты')
    add_para(
        doc,
        'В работе применены актуальные на 2026 год версии библиотек и платформ:'
        ' Python 3.12, Django 4.2 LTS, Django REST Framework 3.14, '
        'djoser 2.2 и djangorestframework-simplejwt 5.3 (аутентификация по JWT), '
        'drf-spectacular 0.27 (генерация OpenAPI и документации Swagger UI / ReDoc), '
        'django-filter 24, Pillow 10 (загрузка изображений), '
        'whitenoise 6 (отдача статики из контейнера). '
        'СУБД — PostgreSQL 16; WSGI-сервер — gunicorn 23; '
        'обратный прокси — nginx 1.25; контейнеризация — Docker Engine 27 '
        'с использованием docker-compose. Тестирование API — Postman 11 '
        'и newman; контроль версий — git (репозиторий на GitHub).'
    )

    add_heading_section(doc, 'Структура работы')
    add_para(
        doc,
        'Работа состоит из введения, пяти глав, заключения, списка использованных '
        'источников и трёх приложений. В первой главе проводится анализ '
        'предметной области и формулируются требования к проектируемой системе. '
        'Во второй главе описано проектирование: архитектура, модели данных, '
        'Use Case диаграммы, API-контракт, права доступа и бизнес-валидации. '
        'В третьей главе изложены детали реализации на Django REST Framework. '
        'В четвёртой главе описано тестирование через Postman-коллекцию '
        'и прогон через newman. В пятой главе раскрыто развёртывание системы '
        'в Docker, приведена Deployment-диаграмма. В заключении обобщены '
        'результаты работы и намечены направления развития.'
    )

    # ╔══════════════════════════════════════════╗
    # ║   ГЛАВА 1. АНАЛИТИКА И ПОСТАНОВКА        ║
    # ╚══════════════════════════════════════════╝
    add_heading_chapter(doc, 'ГЛАВА 1. АНАЛИТИКА И ПОСТАНОВКА ЗАДАЧИ')

    add_heading_section(doc, '1.1. Описание предметной области')
    add_para(
        doc,
        'Базовый Kittygram — это серверная часть социального сервиса для '
        'владельцев котов. Зарегистрированный пользователь может создавать '
        'у себя в профиле карточки своих котов, указывать их имя, окрас, '
        'год рождения и краткое описание, прикреплять фотографию, '
        'присваивать тематические теги (например, «пушистый», «спокойный», '
        '«игривый») и фиксировать «достижения» (например, «Длинная шерсть», '
        '«Чемпион двора», «Громко мурлычет»). Любой посетитель — даже без '
        'регистрации — может посмотреть общий каталог котов и список '
        'достижений.'
    )
    add_para(
        doc,
        'Расширение «мини-чат по заявкам» вводит в предметную область новую '
        'сущность — кото-ивент. Кото-ивент представляет собой событие, '
        'на которое владельцы котов могут собираться вместе со своими '
        'питомцами: фотосессия в парке, прогулка, выставка, тематическая '
        'встреча владельцев конкретной породы. Любой авторизованный '
        'пользователь может стать организатором кото-ивента, указав место, '
        'дату начала и окончания, краткое описание и количество мест. '
        'Другие пользователи могут увидеть ивент в общем списке и подать '
        'заявку на участие, выбрав, с каким именно из своих котов они хотят '
        'прийти, и приложив сопроводительное сообщение.'
    )
    add_para(
        doc,
        'Ключевая особенность сценария — автоматическое открытие диалога '
        'по каждой поданной заявке. Сразу после создания заявки на стороне '
        'сервера через сигнал post_save создаётся пара объектов Dialog '
        'и Message: первое сообщение в диалоге формируется из '
        'сопроводительного текста заявки. Дальше организатор и заявитель '
        'обмениваются сообщениями по REST-эндпоинту, уточняя детали участия '
        '(возраст и темперамент кота, время приезда, особенности '
        'оборудования). Организатор имеет возможность одобрить заявку '
        'или отклонить её через кастомное действие. При финализации заявки '
        '(одобрении, отклонении или отзыве самим заявителем) диалог '
        'автоматически закрывается, и попытка отправить в него новое '
        'сообщение возвращает HTTP 403.'
    )
    add_para(
        doc,
        'Такая модель взаимодействия покрывает реальный пользовательский '
        'сценарий: «организатор события — желающий присоединиться '
        'участник», подразумевает естественный обмен сообщениями для '
        'согласования деталей и обеспечивает корректную бизнес-логику '
        '(например, защиту от подачи заявки на ивент после его начала, '
        'превышения количества мест, дублирующих заявок с одним и тем же '
        'котом, попыток одобрить чужую заявку).'
    )

    add_heading_section(doc, '1.2. Функциональные требования')
    add_para(
        doc,
        'Функциональные требования к системе сгруппированы в таблице 1. '
        'Они охватывают как возможности базового Kittygram, так и расширения '
        '«мини-чат по заявкам».',
    )
    add_table_caption(doc, 1, 'Функциональные требования к системе')
    add_table(
        doc,
        ['Группа', 'Возможность', 'Кому доступно'],
        [
            ['Аутентификация', 'Регистрация по логину/паролю', 'все'],
            ['Аутентификация', 'Получение и обновление JWT', 'зарегистрированные'],
            ['Каталог котов', 'Просмотр списка котов с фильтрами и пагинацией',
             'все, включая гостей'],
            ['Каталог котов', 'Создание/редактирование/удаление кота',
             'владелец карточки'],
            ['Каталог котов', 'Загрузка изображения для кота',
             'владелец карточки'],
            ['Каталог котов', 'Просмотр тегов и достижений',
             'все'],
            ['Кото-ивенты', 'Просмотр списка ивентов', 'все'],
            ['Кото-ивенты', 'Создание/редактирование ивента',
             'организатор'],
            ['Заявки', 'Подача заявки на участие в ивенте со своим котом',
             'зарегистрированные'],
            ['Заявки', 'Просмотр заявок на собственный ивент',
             'организатор'],
            ['Заявки', 'Изменение статуса (approve/reject/cancel)',
             'role-aware'],
            ['Чат по заявке', 'Отправка и чтение сообщений в диалоге',
             'участники'],
            ['Чат по заявке', 'Отметка чужого сообщения прочитанным',
             'не-автор сообщения'],
            ['Администрирование', 'Модерация достижений',
             'администратор'],
        ],
        col_widths_cm=[3, 8, 5],
    )

    add_heading_section(doc, '1.3. Нефункциональные требования')
    add_para(
        doc,
        'К нефункциональным требованиям к системе относятся следующие группы '
        'характеристик.',
    )
    add_heading_subsection(doc, 'Безопасность')
    add_bullet(doc, 'Аутентификация по краткоживущему access-токену JWT '
                    '(60 минут) и обновляющему refresh-токену (7 дней) согласно RFC 7519.')
    add_bullet(doc, 'Все секреты — SECRET_KEY, пароли БД, реквизиты — вынесены '
                    'в .env-файл, который не попадает в репозиторий; в репозитории '
                    'хранится .env.example с заглушками.')
    add_bullet(doc, 'Throttling: 50 запросов в час для анонимных пользователей '
                    'и 1000 запросов в час для авторизованных.')
    add_bullet(doc, 'Объектные права (object-level permissions) на изменение '
                    'и удаление сущностей — только для их владельцев/организаторов.')

    add_heading_subsection(doc, 'Производительность и масштабируемость')
    add_bullet(doc, 'Все списочные эндпоинты по умолчанию возвращают по 10 элементов '
                    'на страницу с навигацией через ?page=, что предотвращает '
                    'перегрузку трафика.')
    add_bullet(doc, 'Gunicorn запускается с тремя worker-процессами, что позволяет '
                    'обрабатывать запросы параллельно даже при отсутствии асинхронности.')
    add_bullet(doc, 'Реверс-прокси nginx отдаёт статические и медиафайлы напрямую, '
                    'минуя Python-приложение, что разгружает backend.')

    add_heading_subsection(doc, 'Сопровождаемость и удобство')
    add_bullet(doc, 'Автоматически генерируемая интерактивная документация API '
                    '(Swagger UI и ReDoc) на отдельных URL.')
    add_bullet(doc, 'Однокомандное развёртывание через docker compose up — '
                    'система поднимается из чистого состояния за минуты.')
    add_bullet(doc, 'Файлы миграций версионируются вместе с кодом — состояние '
                    'схемы БД полностью воспроизводимо.')

    add_heading_subsection(doc, 'Совместимость')
    add_bullet(doc, 'Поддерживаемые ОС: Linux и Windows; единственное требование '
                    '— наличие Docker Engine версии 27 и выше.')
    add_bullet(doc, 'Версии Python: 3.11 и 3.12 (внутри контейнера используется 3.12-slim).')
    add_bullet(doc, 'Браузеры — любые современные с поддержкой стандарта ES2018+ '
                    'для интерактивной Swagger UI.')

    add_heading_section(doc, '1.4. Роли пользователей')
    add_para(
        doc,
        'В системе выделено пять ролей пользователей. Они описывают не отдельные '
        'технические сущности, а скорее «контексты», в которых один и тот же '
        'учётный пользователь может действовать. Например, один и тот же '
        'зарегистрированный пользователь может одновременно быть и организатором '
        '(в одном ивенте), и заявителем (в другом).',
    )
    add_table_caption(doc, 2, 'Роли пользователей и их основные возможности')
    add_table(
        doc,
        ['Роль', 'Описание', 'Ключевые операции'],
        [
            ['Гость', 'Незарегистрированный посетитель',
             'просмотр каталога котов; просмотр списка кото-ивентов'],
            ['Пользователь', 'Зарегистрированный владелец котов',
             'регистрация, вход, CRUD котов, загрузка фото, просмотр тегов и достижений'],
            ['Организатор', 'Зарегистрированный пользователь, создавший хотя бы один ивент',
             'создание ивента, просмотр заявок на свой ивент, одобрение/отклонение, '
             'переписка в диалогах своих ивентов'],
            ['Заявитель', 'Зарегистрированный пользователь, подавший заявку',
             'подача заявки, отзыв своей заявки, переписка в диалоге, '
             'отметка чужого сообщения прочитанным'],
            ['Администратор', 'Сотрудник с правами is_staff/is_superuser',
             'модерация достижений, доступ к Django Admin'],
        ],
        col_widths_cm=[2.5, 6, 7.5],
    )

    add_heading_section(doc, '1.5. Ограничения и допущения')
    add_bullet(doc, 'Чат реализован как REST-эндпоинт без поддержки WebSocket '
                    'и Channels: обновление сообщений на стороне клиента '
                    'предполагается через периодический опрос '
                    '(или Server-Sent Events на следующих итерациях).')
    add_bullet(doc, 'Регистрация — без подтверждения по электронной почте; '
                    'это упрощает локальную разработку и тестирование.')
    add_bullet(doc, 'Уведомления (e-mail, push) не входят в объём текущей работы.')
    add_bullet(doc, 'Геолокация ивентов представлена строкой в свободной форме; '
                    'привязка к карте и поиск по радиусу относятся к перспективам развития.')
    add_bullet(doc, 'Одобрение заявки расходует одно «место» в capacity ивента; '
                    'отклонение/отзыв — освобождают место.')

    # ╔══════════════════════════════════════════╗
    # ║   ГЛАВА 2. ПРОЕКТИРОВАНИЕ                ║
    # ╚══════════════════════════════════════════╝
    add_heading_chapter(doc, 'ГЛАВА 2. ПРОЕКТИРОВАНИЕ')

    add_heading_section(doc, '2.1. Архитектура приложения')
    add_para(
        doc,
        'Приложение спроектировано в соответствии с классической трёхуровневой '
        'архитектурой web-приложения: реверс-прокси (nginx), backend на Python '
        '(Django + DRF под управлением gunicorn) и реляционная СУБД (PostgreSQL). '
        'Все три компонента развёрнуты в виде отдельных Docker-контейнеров '
        'и связаны Docker-сетью. Подробная схема развёртывания приведена '
        'в главе 5.',
    )
    add_para(
        doc,
        'Серверная часть Kittygram декомпозирована на два Django-приложения: '
        'cats (базовая часть — коты, теги, достижения) и events (расширение — '
        'ивенты, заявки, диалоги, сообщения). Такое разделение позволяет '
        'эволюционировать каждый из модулей независимо, а в перспективе — '
        'выделить расширение в отдельный микросервис, если для этого '
        'появится бизнес-обоснование.',
    )

    add_heading_section(doc, '2.2. Use Case диаграммы')
    add_para(
        doc,
        'Для наглядного описания взаимодействия пользователей с системой '
        'построены две диаграммы вариантов использования (Use Case). '
        'Декомпозиция на две диаграммы выбрана для повышения читаемости: '
        'одна охватывает базовую часть Kittygram и администрирование, '
        'другая — расширение «мини-чат по заявкам».',
    )
    add_image(doc, DIAGRAMS / 'usecase_base.png', width_cm=12.0,
              caption='Рисунок 1 — Use Case диаграмма базовой части Kittygram')
    add_para(
        doc,
        'На рисунке 1 показаны акторы базовой части: гость, зарегистрированный '
        'пользователь и администратор. Гость может только просмотреть список '
        'котов; зарегистрированный пользователь может управлять собственными '
        'котами (создание профиля, загрузка фото) и просматривать таксономию '
        '(теги, достижения); администратор отвечает за модерацию достижений.',
    )
    add_image(doc, DIAGRAMS / 'usecase_chat.png', width_cm=12.0,
              caption='Рисунок 2 — Use Case диаграмма расширения «Мини-чат по заявкам»')
    add_para(
        doc,
        'На рисунке 2 представлены акторы расширения: гость (с правом просмотра '
        'списка кото-ивентов), организатор (создатель ивента) и заявитель '
        '(подающий заявку). Use case «Переписка в диалоге» относится '
        'одновременно к организатору и к заявителю, поскольку диалог '
        'разворачивается между ними. Use case «Отметить сообщение прочитанным» '
        'доступен только не-автору сообщения, что обеспечивается '
        'permission-классом и проверкой в сериализаторе.',
    )

    add_heading_section(doc, '2.3. Модели данных')
    add_para(
        doc,
        'Доменные модели разделены на две группы — модели базового Kittygram '
        '(приложение cats) и модели расширения (приложение events). '
        'Их атрибуты, ключевые ограничения и связи приведены в таблицах 3 и 4.',
    )

    add_table_caption(doc, 3, 'Модели базового Kittygram (приложение cats)')
    add_table(
        doc,
        ['Модель', 'Назначение', 'Ключевые поля', 'Связи', 'Ограничения'],
        [
            ['Cat', 'Карточка кота',
             'name, color, birth_year, description, image, owner, created_at',
             'FK owner→User',
             'name 2..32; birth_year 1990..текущий; UNIQUE(owner, name)'],
            ['Achievement', 'Достижение',
             'name', 'M2M через CatAchievement', 'name UNIQUE, 2..64'],
            ['CatAchievement', 'Связь «кот↔достижение»',
             'cat, achievement, achieved_at',
             'FK cat, FK achievement', 'UNIQUE(cat, achievement)'],
            ['Tag', 'Тег',
             'name, slug',
             'M2M cats',
             'slug UNIQUE, regex ^[-a-zA-Z0-9_]+$'],
        ],
        col_widths_cm=[2.5, 3, 4, 3.5, 4],
    )

    add_table_caption(doc, 4, 'Модели расширения «Мини-чат по заявкам» '
                              '(приложение events)')
    add_table(
        doc,
        ['Модель', 'Назначение', 'Ключевые поля', 'Связи', 'Ограничения'],
        [
            ['Event',
             'Кото-ивент с лимитом мест',
             'title, description, location, starts_at, ends_at, capacity',
             'FK organizer→User',
             'title 4..120; capacity 1..100; CHECK ends_at>starts_at; '
             'на create starts_at>now()'],
            ['Application',
             'Заявка на участие в ивенте',
             'message_text, status (pending/approved/rejected/cancelled)',
             'FK event, FK cat, FK applicant→User',
             'UNIQUE(event, cat) — один кот не подаётся дважды'],
            ['Dialog',
             'Диалог по заявке (автогенерация сигналом post_save)',
             'is_closed, created_at',
             'OneToOne application→Application',
             'OneToOne обеспечивает уникальность'],
            ['Message',
             'Сообщение в диалоге',
             'text, is_read, created_at',
             'FK dialog→Dialog, FK author→User',
             'text не пустой; ordering по created_at'],
        ],
        col_widths_cm=[2.5, 3, 3.5, 3, 4.5],
    )

    add_heading_section(doc, '2.4. API-контракт')
    add_para(
        doc,
        'API-контракт представляет собой исчерпывающий перечень всех '
        'HTTP-эндпоинтов системы с указанием URL, метода, назначения, '
        'требуемой аутентификации и возможных кодов ответа. '
        'Полная таблица контракта вынесена в приложение Б; в таблице 5 '
        'ниже приведена сокращённая версия — наиболее значимые эндпоинты.',
    )
    add_table_caption(doc, 5, 'Сокращённый API-контракт (полная версия — приложение Б)')
    add_table(
        doc,
        ['№', 'URL', 'Метод', 'Назначение', 'Auth', 'Коды'],
        [
            ['1', '/api/v1/users/', 'POST', 'Регистрация (Djoser)',
             'public', '201/400'],
            ['2', '/api/v1/auth/jwt/create/', 'POST', 'Получить JWT',
             'public', '200/401'],
            ['3', '/api/v1/cats/', 'GET / POST',
             'Список / создание кота', 'JWT для POST', '200/201/400/401'],
            ['4', '/api/v1/cats/{id}/upload_image/', 'POST',
             'Загрузить фото', 'JWT, владелец', '200/401/403'],
            ['5', '/api/v1/events/', 'GET / POST',
             'Список / создание ивента', 'JWT для POST', '200/201/401'],
            ['6', '/api/v1/events/{id}/applications/', 'GET',
             'Заявки на ивент', 'JWT', '200/401/404'],
            ['7', '/api/v1/applications/', 'POST / GET',
             'Подача заявки / мои заявки', 'JWT', '201/400/401'],
            ['8', '/api/v1/applications/{id}/set_status/', 'POST',
             'Сменить статус (approve/reject/cancel)',
             'JWT, role-aware', '200/400/403/404'],
            ['9', '/api/v1/dialogs/{id}/messages/', 'GET / POST',
             'Чат по заявке', 'JWT, участник', '200/201/401/403'],
            ['10', '/api/v1/messages/{id}/mark_read/', 'POST',
             'Отметить прочитанным', 'JWT, не автор', '200/400/403'],
        ],
        col_widths_cm=[1, 4.2, 2.4, 3.6, 2.5, 1.8],
    )

    add_heading_section(doc, '2.5. Права доступа')
    add_para(
        doc,
        'Права доступа реализованы на двух уровнях: уровень класса представления '
        '(permission_classes из DRF) и уровень queryset (фильтрация выдачи '
        'в зависимости от пользователя). Сводка прав приведена в таблице 6.',
    )
    add_table_caption(doc, 6, 'Права доступа к ключевым эндпоинтам')
    add_table(
        doc,
        ['Эндпоинт', 'Чтение', 'Изменение', 'Permission-класс'],
        [
            ['/cats/', 'public', 'JWT; владелец',
             'IsAuthenticatedOrReadOnly + IsOwnerOrReadOnly'],
            ['/events/', 'public', 'JWT; организатор для PATCH/DELETE',
             'IsOrganizerOrReadOnly'],
            ['/applications/', 'JWT; queryset фильтрует по applicant и event.organizer',
             'JWT', 'IsAuthenticated'],
            ['/applications/{id}/set_status/', '—',
             'approve/reject — организатор; cancel — заявитель',
             'IsAuthenticated + IsApplicationParticipant + role-aware валидация'],
            ['/dialogs/{id}/messages/', 'JWT, участник',
             'JWT, участник, диалог открыт',
             'IsAuthenticated + IsDialogParticipant'],
            ['/messages/{id}/mark_read/', '—',
             'участник диалога, не автор сообщения',
             'IsAuthenticated + IsDialogParticipant + author≠user'],
        ],
        col_widths_cm=[4, 3, 4, 5],
    )

    add_heading_section(doc, '2.6. Бизнес-валидации')
    add_para(
        doc,
        'Помимо проверок на уровне БД (UNIQUE, CHECK), на уровне сериализаторов '
        'реализован набор бизнес-валидаций, отражающих правила предметной области. '
        'Сводка приведена в таблице 7.',
    )
    add_table_caption(doc, 7, 'Бизнес-валидации сценария «Мини-чат по заявкам»')
    add_table(
        doc,
        ['Правило', 'Где проверяется', 'Ответ при нарушении'],
        [
            ['Нельзя подать заявку на собственный ивент',
             'ApplicationSerializer.validate',
             '400 — {"event": ["Вы организатор этого ивента — заявку подавать не нужно."]}'],
            ['Нельзя подать заявку с чужим котом',
             'ApplicationSerializer.validate',
             '400 — {"cat": ["Можно подать заявку только со своим котом."]}'],
            ['Нельзя подать заявку на прошедший/начавшийся ивент',
             'ApplicationSerializer.validate',
             '400 — {"event": ["Ивент уже начался — заявку подавать поздно."]}'],
            ['Один кот = одна заявка на ивент',
             'Application.Meta.constraints + serializer pre-check',
             '400 — {"cat": ["Заявка с этим котом на этот ивент уже подана."]}'],
            ['Capacity не превышается при approve',
             'ApplicationStatusSerializer.validate_status',
             '400 — ["На ивенте больше нет свободных мест."]'],
            ['Сменить статус: approve/reject — организатор; cancel — заявитель',
             'ApplicationStatusSerializer (role-aware)',
             '400/403 с пояснением'],
            ['Писать в закрытый диалог нельзя',
             'DialogMessagesViewSet.create',
             '403 — {"detail": "Диалог закрыт — писать в него нельзя."}'],
        ],
        col_widths_cm=[5, 4.5, 6.5],
    )

    # ╔══════════════════════════════════════════╗
    # ║   ГЛАВА 3. РЕАЛИЗАЦИЯ ПРОЕКТА            ║
    # ╚══════════════════════════════════════════╝
    add_heading_chapter(doc, 'ГЛАВА 3. РЕАЛИЗАЦИЯ ПРОЕКТА')

    add_heading_section(doc, '3.1. Структура проекта')
    add_para(
        doc,
        'Дерево каталогов репозитория проекта показано ниже. Корневой каталог '
        'kitty/ содержит конфигурации Docker, файл зависимостей requirements.txt, '
        'модули settings-пакета kittygram/ и два прикладных Django-приложения: '
        'cats/ и events/.',
    )
    add_code(
        doc,
        'kitty/\n'
        '├── .env.example            конфигурация (заглушки)\n'
        '├── README.md               инструкция по запуску\n'
        '├── manage.py\n'
        '├── requirements.txt        зависимости Python\n'
        '├── Dockerfile              сборка backend-образа\n'
        '├── docker-compose.yml      db + backend + nginx\n'
        '├── nginx/default.conf      reverse-proxy конфиг\n'
        '├── kittygram/              settings-пакет (settings.py, urls.py, wsgi.py)\n'
        '├── cats/                   приложение базового Kittygram\n'
        '├── events/                 приложение мини-чата по заявкам\n'
        '├── core/                   общие утилиты (pagination, permissions)\n'
        '├── postman/                Postman-коллекции v2.1\n'
        '├── scripts/                seed.py, build_*_report.py\n'
        '└── docs/                   диаграммы и скриншоты\n',
    )

    add_heading_section(doc, '3.2. Стек реализации')
    add_para(
        doc,
        'Перечень используемых пакетов Python и их назначение приведены в таблице 8.',
    )
    add_table_caption(doc, 8, 'Стек реализации')
    add_table(
        doc,
        ['Пакет', 'Версия', 'Назначение'],
        [
            ['Django', '4.2.16 LTS', 'Веб-фреймворк'],
            ['djangorestframework', '3.14.0', 'Реализация REST API'],
            ['djoser', '2.2.3', 'Готовые эндпоинты регистрации/входа/восстановления'],
            ['djangorestframework-simplejwt', '5.3.1', 'JWT токены (access/refresh)'],
            ['drf-spectacular', '0.27.2', 'Генерация OpenAPI схемы и Swagger UI/ReDoc'],
            ['django-filter', '24.3', 'Фильтрация по полям модели'],
            ['Pillow', '10.4.0', 'Работа с изображениями (загрузка фото кота)'],
            ['psycopg2-binary', '2.9.9', 'Драйвер PostgreSQL'],
            ['gunicorn', '23.0.0', 'WSGI-сервер'],
            ['python-dotenv', '1.0.1', 'Чтение .env-файла'],
            ['whitenoise', '6.7.0', 'Отдача статики из контейнера'],
        ],
        col_widths_cm=[5, 2.5, 8],
    )

    add_heading_section(doc, '3.3. Реализация моделей')
    add_para(
        doc,
        'Модели объявлены в cats/models.py и events/models.py. Ниже приведён '
        'фрагмент реализации модели Event с полным набором ограничений.',
    )
    add_code(
        doc,
        'class Event(models.Model):\n'
        '    organizer = models.ForeignKey(\n'
        '        User, on_delete=models.CASCADE, related_name=\'events_organized\'\n'
        '    )\n'
        '    title = models.CharField(max_length=120, validators=[MinLengthValidator(4)])\n'
        '    description = models.TextField(blank=True)\n'
        '    location = models.CharField(max_length=255)\n'
        '    starts_at = models.DateTimeField()\n'
        '    ends_at = models.DateTimeField()\n'
        '    capacity = models.PositiveSmallIntegerField(\n'
        '        validators=[MinValueValidator(1), MaxValueValidator(100)]\n'
        '    )\n'
        '    created_at = models.DateTimeField(auto_now_add=True)\n'
        '\n'
        '    class Meta:\n'
        '        ordering = (\'-starts_at\',)\n'
        '        constraints = [\n'
        '            models.CheckConstraint(\n'
        '                check=Q(ends_at__gt=F(\'starts_at\')),\n'
        '                name=\'event_ends_after_starts\',\n'
        '            ),\n'
        '        ]',
    )

    add_heading_section(doc, '3.4. Реализация ViewSet‘ов')
    add_para(
        doc,
        'API реализован в виде набора ViewSet’ов. Базовая часть — '
        'CatViewSet, AchievementViewSet, TagViewSet — в cats/views.py. '
        'Расширение — EventViewSet, ApplicationViewSet, DialogMessagesViewSet, '
        'MessageActionViewSet — в events/views.py. Каждый ViewSet наследуется '
        'от ModelViewSet (или от выборки нужных mixins) и регистрируется '
        'в DefaultRouter в файле urls.py соответствующего приложения. '
        'Кастомные действия (action) подключены через декоратор @action '
        'из rest_framework.decorators.',
    )
    add_para(
        doc,
        'Пример объявления кастомного действия set_status, изменяющего '
        'статус заявки и автоматически закрывающего диалог при '
        'переходе в финальное состояние:',
    )
    add_code(
        doc,
        '@action(detail=True, methods=[\'post\'],\n'
        '        permission_classes=[IsAuthenticated, IsApplicationParticipant])\n'
        'def set_status(self, request, pk=None):\n'
        '    application = self.get_object()\n'
        '    serializer = ApplicationStatusSerializer(\n'
        '        data=request.data, context={\'request\': request,\n'
        '                                    \'application\': application},\n'
        '    )\n'
        '    serializer.is_valid(raise_exception=True)\n'
        '    application.status = serializer.validated_data[\'status\']\n'
        '    application.save(update_fields=(\'status\', \'updated_at\'))\n'
        '    if application.status in {\n'
        '        Application.Status.APPROVED, Application.Status.REJECTED,\n'
        '        Application.Status.CANCELLED,\n'
        '    }:\n'
        '        application.dialog.is_closed = True\n'
        '        application.dialog.save(update_fields=(\'is_closed\',))\n'
        '    return Response(ApplicationSerializer(application).data)',
    )

    add_heading_section(doc, '3.5. Сериализаторы')
    add_para(
        doc,
        'Сериализаторы (DRF Serializer) выполняют две задачи: преобразование '
        'между Python-объектами и JSON, а также валидацию входящих данных. '
        'Все бизнес-валидации сценария «мини-чат по заявкам», описанные '
        'в разделе 2.6, реализованы именно на уровне сериализаторов '
        '(events/serializers.py). Например, проверка «нельзя подать заявку '
        'на собственный ивент» оформлена как метод validate(self, attrs), '
        'который сравнивает event.organizer с self.context["request"].user '
        'и при совпадении генерирует ValidationError. ApplicationStatusSerializer '
        'отдельно реализует role-aware логику: метод validate_status проверяет, '
        'что текущий пользователь имеет право на запрашиваемый переход '
        '(approve/reject — организатор; cancel — заявитель) и не превышает '
        'лимит capacity при approve.',
    )

    add_heading_section(doc, '3.6. Permissions')
    add_para(
        doc,
        'Кастомные классы прав доступа размещены в core/permissions.py '
        '(IsOwnerOrReadOnly) и events/permissions.py '
        '(IsOrganizerOrReadOnly, IsApplicationParticipant, IsDialogParticipant). '
        'Каждый класс реализует метод has_object_permission, проверяющий '
        'отношение между текущим пользователем и объектом запроса.',
    )

    add_heading_section(doc, '3.7. Сигналы автосоздания диалога')
    add_para(
        doc,
        'Автоматическое открытие диалога по каждой поданной заявке '
        'реализовано через механизм сигналов Django. В events/signals.py '
        'функция-обработчик подписана на сигнал post_save модели Application: '
        'при создании новой заявки она создаёт связанный Dialog '
        'и помещает в него первое сообщение Message с содержимым '
        'message_text заявки. Это позволяет сохранить простоту API '
        '(клиенту не нужно отдельно создавать диалог) и гарантирует '
        'консистентность: заявка без диалога невозможна.',
    )

    add_heading_section(doc, '3.8. Фильтрация, поиск, пагинация')
    add_para(
        doc,
        'Списочные эндпоинты поддерживают комбинацию из трёх механизмов '
        'выборки: фильтрацию по точным значениям полей через django-filter, '
        'полнотекстовый поиск по нескольким полям через DRF SearchFilter '
        'и сортировку через OrderingFilter. Пагинация реализована классом '
        'core.pagination.DefaultPagination (PageNumberPagination, page_size=10, '
        'max_page_size=100). Например, эндпоинт /api/v1/events/ принимает '
        'параметры ?starts_after=2026-05-01, ?organizer=123, ?search=парк, '
        '?ordering=-starts_at, ?page=2.',
    )

    add_heading_section(doc, '3.9. Обработка ошибок')
    add_para(
        doc,
        'Все ошибки клиента возвращаются в формате JSON с понятным русскоязычным '
        'описанием. Стандартный обработчик исключений DRF возвращает 400 при '
        'ошибках валидации, 401 при отсутствии аутентификации, 403 — при '
        'отсутствии прав, 404 — для несуществующих объектов. Объектные права '
        'на просмотр (queryset-фильтрация) маскируют чужие объекты как 404, '
        'что соответствует рекомендуемой практике (раскрывать факт '
        'существования объекта только тем, кто имеет право его видеть).',
    )

    add_heading_section(doc, '3.10. Документация API')
    add_para(
        doc,
        'Документация генерируется автоматически из аннотаций сериализаторов '
        'и описаний действий ViewSet’ов через библиотеку drf-spectacular. '
        'Доступны три представления: '
        'JSON-схема OpenAPI 3.0 на /api/schema/, '
        'интерактивная документация Swagger UI на /api/schema/swagger-ui/ '
        'и документация ReDoc на /api/schema/redoc/. '
        'Скриншоты Swagger UI и ReDoc приведены в главе 4.',
    )

    # ╔══════════════════════════════════════════╗
    # ║   ГЛАВА 4. ТЕСТИРОВАНИЕ                  ║
    # ╚══════════════════════════════════════════╝
    add_heading_chapter(doc, 'ГЛАВА 4. ТЕСТИРОВАНИЕ И ПРОВЕРКА РАБОТОСПОСОБНОСТИ')

    add_heading_section(doc, '4.1. Подход к тестированию')
    add_para(
        doc,
        'Для проверки работоспособности API применён подход «чёрного ящика»: '
        'через HTTP-клиент Postman собирается коллекция запросов, '
        'имитирующих типовые пользовательские сценарии. Каждый запрос '
        'дополняется тест-скриптом на JavaScript, который проверяет '
        'код ответа, структуру тела и сохраняет извлечённые значения '
        '(идентификаторы, токены) в переменные коллекции для последующих '
        'запросов. Для воспроизводимого пакетного прогона используется '
        'утилита newman — CLI-runner от Postman. Newman запускается '
        'с reporter’ом htmlextra и формирует HTML-отчёт о прохождении '
        'всех запросов и тестов.',
    )

    add_heading_section(doc, '4.2. Postman-коллекция')
    add_para(
        doc,
        'Postman-коллекция расширения «мини-чат по заявкам» содержит 17 запросов, '
        'разбитых на четыре функциональные группы: «Auth & seed» (регистрация и '
        'получение токенов для тестовых пользователей), «Events» (CRUD ивентов), '
        '«Applications & dialogs» (подача заявки, смена статуса, переписка) '
        'и «Negative» (негативные сценарии: чужой кот, превышение capacity, '
        'попытка писать в закрытый диалог). Файл коллекции — '
        'kittygram_creative.postman_collection.json — приведён '
        'в приложении А и также доступен в каталоге postman/ репозитория.',
    )

    add_heading_section(doc, '4.3. Положительные сценарии')
    add_para(
        doc,
        'Основной положительный сценарий покрывает полный happy-path: '
        'регистрация двух пользователей (alice и bob), получение JWT, '
        'создание ивента alice’ой, подача заявки bob’ом со своим '
        'котом Мурзиком, переписка в диалоге, одобрение заявки alice’ой, '
        'попытка bob’а написать в закрытый диалог (HTTP 403). '
        'Скриншоты Swagger UI и ReDoc, на которых видны соответствующие '
        'эндпоинты, приведены ниже.',
    )
    add_image(doc, SCREENS / '01_swagger_top.png', width_cm=15.0,
              caption='Рисунок 3 — Swagger UI: верхняя часть API-документации')
    add_image(doc, SCREENS / '02_swagger_dialogs_events.png', width_cm=15.0,
              caption='Рисунок 4 — Swagger UI: эндпоинты dialogs и events')
    add_image(doc, SCREENS / '03_swagger_application_post.png', width_cm=15.0,
              caption='Рисунок 5 — Swagger UI: интерактивный POST /applications/')
    add_image(doc, SCREENS / '04_redoc_events_create.png', width_cm=15.0,
              caption='Рисунок 6 — ReDoc: спецификация POST /events/')
    add_image(doc, SCREENS / '05_redoc_set_status.png', width_cm=15.0,
              caption='Рисунок 7 — ReDoc: спецификация кастомного '
                      'действия set_status')

    add_heading_section(doc, '4.4. Негативные сценарии')
    add_para(
        doc,
        'В коллекцию включены негативные сценарии, проверяющие корректность '
        'бизнес-валидаций: подача заявки с чужим котом возвращает HTTP 400 '
        'с понятным сообщением об ошибке; попытка одобрить заявку чужим '
        'пользователем возвращает HTTP 403; попытка подать вторую заявку '
        'с тем же котом — HTTP 400; попытка писать в закрытый диалог '
        'после approve — HTTP 403. Все эти проверки выполняются как '
        'в Postman-тестах, так и в Django Admin (рисунки 8–11) — для контроля '
        'состояния БД после прогона.',
    )
    add_image(doc, SCREENS / '06_admin_index.png', width_cm=15.0,
              caption='Рисунок 8 — Django Admin: общий вид')
    add_image(doc, SCREENS / '07_admin_events.png', width_cm=15.0,
              caption='Рисунок 9 — Django Admin: список кото-ивентов после прогона')
    add_image(doc, SCREENS / '08_admin_applications.png', width_cm=15.0,
              caption='Рисунок 10 — Django Admin: заявки и их статусы')
    add_image(doc, SCREENS / '09_admin_messages.png', width_cm=15.0,
              caption='Рисунок 11 — Django Admin: сообщения в диалогах')

    add_heading_section(doc, '4.5. Прогон через newman')
    add_para(
        doc,
        'Все запросы коллекции прогоняются командой '
        '«newman run kittygram_creative.postman_collection.json -r htmlextra». '
        'На рисунках 12–14 показаны фрагменты HTML-отчёта newman: '
        'верхняя сводка, средняя часть с детальным разбором каждого запроса '
        'и итоговая статистика прохождения тестов.',
    )
    add_image(doc, SCREENS / '10_newman_top.png', width_cm=15.0,
              caption='Рисунок 12 — Newman: верхняя сводка прогона')
    add_image(doc, SCREENS / '11_newman_middle.png', width_cm=15.0,
              caption='Рисунок 13 — Newman: разбор отдельных запросов')
    add_image(doc, SCREENS / '12_newman_bottom.png', width_cm=15.0,
              caption='Рисунок 14 — Newman: итоговая статистика прохождения')

    # ╔══════════════════════════════════════════╗
    # ║   ГЛАВА 5. РАЗВЁРТЫВАНИЕ                 ║
    # ╚══════════════════════════════════════════╝
    add_heading_chapter(doc, 'ГЛАВА 5. РАЗВЁРТЫВАНИЕ ПРОЕКТА')

    add_heading_section(doc, '5.1. Контейнеризация (Dockerfile)')
    add_para(
        doc,
        'Backend-приложение упаковано в Docker-образ. В качестве базы выбран '
        'образ python:3.12-slim — лёгкий, актуальный и совместимый '
        'со всеми зависимостями. Команда CMD запускает gunicorn с тремя '
        'worker-процессами, привязанным к 0.0.0.0:8000.',
    )
    add_code(
        doc,
        'FROM python:3.12-slim\n'
        'ENV PYTHONUNBUFFERED=1 PYTHONDONTWRITEBYTECODE=1\n'
        'WORKDIR /app\n'
        'COPY requirements.txt .\n'
        'RUN pip install --no-cache-dir -r requirements.txt\n'
        'COPY . .\n'
        'EXPOSE 8000\n'
        'CMD ["gunicorn", "kittygram.wsgi:application",\n'
        '     "--bind", "0.0.0.0:8000", "--workers", "3"]',
    )

    add_heading_section(doc, '5.2. Многоконтейнерная архитектура (docker-compose)')
    add_para(
        doc,
        'Файл docker-compose.yml описывает три сервиса. Сервис db поднимает '
        'PostgreSQL 16 (образ postgres:16-alpine) и сохраняет данные в томе '
        'pg_data. Сервис backend собирается из локального Dockerfile, '
        'выполняет миграции и collectstatic при старте через входной '
        'shell-скрипт и слушает порт 8000. Сервис nginx (образ '
        'nginx:1.25-alpine) запускается под профилем with-nginx — то есть '
        'опционально (по команде «docker compose --profile with-nginx up»). '
        'Это сделано для того, чтобы локальная разработка могла '
        'осуществляться без nginx (запросы идут напрямую на :8000), '
        'а production-режим — с nginx (на :80 с проксированием).',
    )

    add_heading_section(doc, '5.3. Сетевое взаимодействие')
    add_para(
        doc,
        'Все три контейнера подключены к одной сети kitty_default, '
        'создаваемой docker-compose автоматически. Внутри сети контейнеры '
        'находят друг друга по именам сервисов (db, backend, nginx). '
        'Backend подключается к PostgreSQL по внутренней связи db:5432; '
        'nginx проксирует входящие HTTP-запросы из внешнего мира '
        'на backend:8000.',
    )

    add_heading_section(doc, '5.4. Хранение данных (volumes)')
    add_para(
        doc,
        'Долгоживущие данные хранятся в трёх именованных томах Docker. '
        'Том pg_data содержит файлы PostgreSQL, том static_volume — '
        'статические файлы приложения, собранные через collectstatic, '
        'том media_volume — загруженные пользователями изображения. '
        'Тома static_volume и media_volume монтируются и в backend, '
        'и в nginx, что позволяет nginx отдавать файлы напрямую, '
        'разгружая Python-приложение.',
    )

    add_heading_section(doc, '5.5. Конфигурация (.env)')
    add_para(
        doc,
        'Все секреты и переменные окружения вынесены в файл .env, '
        'расположенный в корне репозитория и исключённый из git через '
        '.gitignore. В репозитории присутствует .env.example с заглушками — '
        'студент копирует его в .env и подставляет свои значения. Ниже '
        'приведено содержимое .env.example.',
    )
    add_code(
        doc,
        'SECRET_KEY=replace-me-with-a-long-random-string\n'
        'DEBUG=False\n'
        'ALLOWED_HOSTS=localhost,127.0.0.1\n'
        'POSTGRES_DB=kittygram\n'
        'POSTGRES_USER=kittygram\n'
        'POSTGRES_PASSWORD=replace-me\n'
        'DB_HOST=db\n'
        'DB_PORT=5432\n',
    )

    add_heading_section(doc, '5.6. Deployment диаграмма')
    add_para(
        doc,
        'Полная схема развёртывания приложения приведена на рисунке 15. '
        'На диаграмме показаны: внешний клиент (браузер, Postman или curl), '
        'хост-машина с Docker Engine, Docker-сеть kitty_default '
        'с тремя контейнерами (kittygram_nginx, kittygram_backend, '
        'kittygram_db), три Docker-тома (static_volume, media_volume, '
        'pg_data) и конфигурационный файл .env, который монтируется '
        'в backend и в db через директиву env_file.',
    )
    add_image(doc, DIAGRAMS / 'deployment.png', width_cm=15.5,
              caption='Рисунок 15 — Deployment-диаграмма: '
                      'развёртывание Kittygram под docker-compose')

    add_heading_section(doc, '5.7. Команды развёртывания')
    add_para(
        doc,
        'Полное развёртывание приложения с нуля выполняется одной командой:',
    )
    add_code(
        doc,
        '# Локальная разработка (без nginx, прямое обращение на :8000)\n'
        'docker compose up -d --build\n'
        '\n'
        '# Production-режим (с nginx на :80)\n'
        'docker compose --profile with-nginx up -d --build\n'
        '\n'
        '# Применение миграций (выполняется автоматически entrypoint‘ом)\n'
        'docker compose exec backend python manage.py migrate\n'
        '\n'
        '# Загрузка демо-данных для скриншотов\n'
        'docker compose exec backend python manage.py shell < scripts/seed.py\n',
    )

    # ╔══════════════════════════════════════════╗
    # ║   ЗАКЛЮЧЕНИЕ                             ║
    # ╚══════════════════════════════════════════╝
    add_heading_chapter(doc, 'ЗАКЛЮЧЕНИЕ')

    add_para(
        doc,
        'В ходе выполнения курсовой работы была спроектирована и реализована '
        'серверная часть проекта Kittygram с поддержкой пользовательского '
        'сценария «мини-чат по заявкам на кото-ивенты». Все поставленные '
        'задачи выполнены в полном объёме.',
    )
    add_para(
        doc,
        'Проведён анализ предметной области; выделены пять ролей пользователей; '
        'сформулированы функциональные и нефункциональные требования. '
        'Спроектированы две группы моделей данных (cats и events), построены '
        'две Use Case диаграммы — для базовой части и для расширения, — '
        'описан полный API-контракт, права доступа и набор бизнес-валидаций. '
        'Разработка велась на актуальном стеке: Python 3.12, Django 4.2 LTS, '
        'Django REST Framework 3.14, PostgreSQL 16, Docker Engine 27.',
    )
    add_para(
        doc,
        'Реализованы 19 эндпоинтов (10 в базовой части, 9 — в расширении) '
        'с двумя кастомными действиями (set_status, mark_read), фильтрацией, '
        'поиском, пагинацией, JWT-аутентификацией и автоматической '
        'OpenAPI-документацией. Бизнес-логика расширения покрыта семью '
        'валидациями, обеспечивающими корректность сценария: проверка '
        'своего ивента, чужого кота, прошедшей даты, дублирующих заявок, '
        'переполнения capacity, role-aware смены статуса, закрытого диалога. '
        'Автоматическое открытие диалога по сигналу post_save обеспечивает '
        'инвариант «у каждой заявки есть диалог».',
    )
    add_para(
        doc,
        'Проведено функциональное тестирование через Postman-коллекцию '
        'из 17 запросов, охватывающих как положительные, так и негативные '
        'сценарии. Прогон через newman формирует HTML-отчёт и подтверждает '
        '100% прохождение тестов. Развёртывание системы реализовано через '
        'docker-compose с тремя сервисами (db, backend, nginx) и тремя '
        'постоянными томами; обеспечена однокомандная установка из чистого '
        'состояния.',
    )

    add_heading_section(doc, 'Возможные направления развития')
    add_bullet(doc, 'Перевод чата на WebSocket / Server-Sent Events через '
                    'Django Channels — для мгновенной доставки сообщений '
                    'без опроса.')
    add_bullet(doc, 'Уведомления — e-mail и push (Web Push API) — '
                    'при поступлении новой заявки и нового сообщения в диалоге.')
    add_bullet(doc, 'Интеграция геолокации: координаты ивентов и поиск '
                    'по радиусу через PostGIS.')
    add_bullet(doc, 'Кеширование частых выборок (например, списка ивентов '
                    'и тегов) через Redis.')
    add_bullet(doc, 'Перенос фотографий котов в S3-совместимое хранилище '
                    '(MinIO, AWS S3) с генерацией миниатюр асинхронно через Celery.')
    add_bullet(doc, 'Мобильный клиент (React Native или Flutter), '
                    'использующий тот же REST API.')
    add_bullet(doc, 'CI/CD: автоматический прогон Postman-коллекции через '
                    'newman в GitHub Actions при каждом push.')

    # ╔══════════════════════════════════════════╗
    # ║   СПИСОК ИСТОЧНИКОВ                      ║
    # ╚══════════════════════════════════════════╝
    add_heading_chapter(doc, 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ')

    add_para(
        doc,
        'Список оформлен в соответствии с требованиями ГОСТ Р 7.0.5-2008 '
        '«Библиографическая ссылка. Общие требования и правила составления». '
        'Все URL проверены и являются кликабельными.',
        indent=False,
    )

    sources = [
        # Документация фреймворков
        [
            'Django Documentation. Version 4.2 LTS [Электронный ресурс]. — URL: ',
            ('link', 'https://docs.djangoproject.com/en/4.2/',
             'https://docs.djangoproject.com/en/4.2/'),
            f' (дата обращения: {ACCESS_DATE}).',
        ],
        [
            'Django REST Framework Documentation [Электронный ресурс]. — URL: ',
            ('link', 'https://www.django-rest-framework.org/',
             'https://www.django-rest-framework.org/'),
            f' (дата обращения: {ACCESS_DATE}).',
        ],
        [
            'Djoser. REST implementation of Django authentication system '
            '[Электронный ресурс]. — URL: ',
            ('link', 'https://djoser.readthedocs.io/en/latest/',
             'https://djoser.readthedocs.io/en/latest/'),
            f' (дата обращения: {ACCESS_DATE}).',
        ],
        [
            'Simple JWT. A JSON Web Token authentication plugin for the Django REST Framework '
            '[Электронный ресурс]. — URL: ',
            ('link', 'https://django-rest-framework-simplejwt.readthedocs.io/en/latest/',
             'https://django-rest-framework-simplejwt.readthedocs.io/en/latest/'),
            f' (дата обращения: {ACCESS_DATE}).',
        ],
        [
            'drf-spectacular. Sane and flexible OpenAPI 3.0 schema generation '
            'for Django REST framework [Электронный ресурс]. — URL: ',
            ('link', 'https://drf-spectacular.readthedocs.io/en/latest/',
             'https://drf-spectacular.readthedocs.io/en/latest/'),
            f' (дата обращения: {ACCESS_DATE}).',
        ],
        [
            'django-filter Documentation [Электронный ресурс]. — URL: ',
            ('link', 'https://django-filter.readthedocs.io/en/stable/',
             'https://django-filter.readthedocs.io/en/stable/'),
            f' (дата обращения: {ACCESS_DATE}).',
        ],
        [
            'Pillow. Python Imaging Library Fork Documentation [Электронный ресурс]. — URL: ',
            ('link', 'https://pillow.readthedocs.io/en/stable/',
             'https://pillow.readthedocs.io/en/stable/'),
            f' (дата обращения: {ACCESS_DATE}).',
        ],
        # Инфраструктура
        [
            'Docker Documentation [Электронный ресурс]. — URL: ',
            ('link', 'https://docs.docker.com/',
             'https://docs.docker.com/'),
            f' (дата обращения: {ACCESS_DATE}).',
        ],
        [
            'Docker Compose Specification [Электронный ресурс]. — URL: ',
            ('link', 'https://docs.docker.com/compose/',
             'https://docs.docker.com/compose/'),
            f' (дата обращения: {ACCESS_DATE}).',
        ],
        [
            'PostgreSQL 16 Documentation [Электронный ресурс]. — URL: ',
            ('link', 'https://www.postgresql.org/docs/16/',
             'https://www.postgresql.org/docs/16/'),
            f' (дата обращения: {ACCESS_DATE}).',
        ],
        [
            'Nginx Documentation [Электронный ресурс]. — URL: ',
            ('link', 'https://nginx.org/en/docs/',
             'https://nginx.org/en/docs/'),
            f' (дата обращения: {ACCESS_DATE}).',
        ],
        [
            'Gunicorn — Python WSGI HTTP Server [Электронный ресурс]. — URL: ',
            ('link', 'https://docs.gunicorn.org/en/stable/',
             'https://docs.gunicorn.org/en/stable/'),
            f' (дата обращения: {ACCESS_DATE}).',
        ],
        # Стандарты и спецификации
        [
            'Jones M., Bradley J., Sakimura N. RFC 7519: JSON Web Token (JWT) '
            '[Электронный ресурс]. — IETF, 2015. — URL: ',
            ('link', 'https://datatracker.ietf.org/doc/html/rfc7519',
             'https://datatracker.ietf.org/doc/html/rfc7519'),
            f' (дата обращения: {ACCESS_DATE}).',
        ],
        [
            'Fielding R., Reschke J. RFC 7231: Hypertext Transfer Protocol (HTTP/1.1): '
            'Semantics and Content [Электронный ресурс]. — IETF, 2014. — URL: ',
            ('link', 'https://datatracker.ietf.org/doc/html/rfc7231',
             'https://datatracker.ietf.org/doc/html/rfc7231'),
            f' (дата обращения: {ACCESS_DATE}).',
        ],
        [
            'Fielding R. T. Architectural Styles and the Design of Network-based '
            'Software Architectures: PhD dissertation [Электронный ресурс]. — '
            'University of California, Irvine, 2000. — URL: ',
            ('link', 'https://www.ics.uci.edu/~fielding/pubs/dissertation/top.htm',
             'https://www.ics.uci.edu/~fielding/pubs/dissertation/top.htm'),
            f' (дата обращения: {ACCESS_DATE}).',
        ],
        [
            'OpenAPI Specification, Version 3.1.0 [Электронный ресурс]. — URL: ',
            ('link', 'https://spec.openapis.org/oas/v3.1.0',
             'https://spec.openapis.org/oas/v3.1.0'),
            f' (дата обращения: {ACCESS_DATE}).',
        ],
        # ГОСТ
        [
            'ГОСТ Р 7.0.5-2008. Система стандартов по информации, библиотечному и '
            'издательскому делу. Библиографическая ссылка. Общие требования и правила '
            'составления [Электронный ресурс]. — М.: Стандартинформ, 2008. — URL: ',
            ('link', 'https://docs.cntd.ru/document/1200063713',
             'https://docs.cntd.ru/document/1200063713'),
            f' (дата обращения: {ACCESS_DATE}).',
        ],
        # Печатные книги
        [
            'Greenfeld D. R., Roy-Greenfeld A. Two Scoops of Django 3.x: Best Practices '
            'for the Django Web Framework. — Two Scoops Press, 2020. — 532 p.',
        ],
        [
            'Richardson L., Amundsen M., Ruby S. RESTful Web APIs. — Sebastopol, CA: '
            'O’Reilly Media, 2013. — 406 p.',
        ],
        [
            'Percival H., Gregory B. Architecture Patterns with Python: Enabling '
            'Test-Driven Development, Domain-Driven Design, and Event-Driven '
            'Microservices. — Sebastopol, CA: O’Reilly Media, 2020. — 304 p.',
        ],
    ]
    for i, parts in enumerate(sources, start=1):
        add_source(doc, i, parts)

    # ╔══════════════════════════════════════════╗
    # ║   ПРИЛОЖЕНИЕ А — Postman                 ║
    # ╚══════════════════════════════════════════╝
    add_heading_chapter(doc, 'ПРИЛОЖЕНИЕ А')
    add_heading_centered(doc, 'Postman-коллекция')
    add_para(
        doc,
        'Полная Postman-коллекция расширения «мини-чат по заявкам» содержится '
        'в файле kittygram_creative.postman_collection.json в каталоге postman/ '
        'репозитория проекта. Коллекция версии Postman v2.1, общий объём — '
        '17 запросов в четырёх группах. Ниже приведён сокращённый перечень '
        'запросов с их ролью в сценарии тестирования.',
    )
    add_table(
        doc,
        ['Группа', '№', 'Запрос', 'Ожидаемый результат'],
        [
            ['Auth & seed', '1', 'POST /users/ (alice)', '201 Created'],
            ['Auth & seed', '2', 'POST /users/ (bob)', '201 Created'],
            ['Auth & seed', '3', 'POST /auth/jwt/create/ (alice)', '200 + токены'],
            ['Auth & seed', '4', 'POST /auth/jwt/create/ (bob)', '200 + токены'],
            ['Events', '5', 'POST /events/ (alice)', '201 Created'],
            ['Events', '6', 'GET /events/', '200 + список из 1 ивента'],
            ['Events', '7', 'POST /cats/ (bob)', '201 Created'],
            ['Apps & dialogs', '8', 'POST /applications/ (bob с котом bob‘а)',
             '201 + Dialog + Message'],
            ['Apps & dialogs', '9', 'GET /events/{id}/applications/ (alice)',
             '200 + 1 заявка'],
            ['Apps & dialogs', '10', 'POST /dialogs/{id}/messages/ (alice)',
             '201 — ответ организатора'],
            ['Apps & dialogs', '11', 'GET /dialogs/{id}/messages/ (bob)',
             '200 + 2 сообщения'],
            ['Apps & dialogs', '12', 'POST /messages/{id}/mark_read/ (bob)',
             '200 — отметка прочитанным'],
            ['Apps & dialogs', '13', 'POST /applications/{id}/set_status/ approved',
             '200 + dialog.is_closed=True'],
            ['Negative', '14', 'POST /applications/ (bob с котом alice)',
             '400 — чужой кот'],
            ['Negative', '15', 'POST /dialogs/{id}/messages/ после approve',
             '403 — диалог закрыт'],
            ['Negative', '16', 'POST /applications/ (alice на свой ивент)',
             '400 — свой ивент'],
            ['Negative', '17', 'POST /applications/{id}/set_status/ approved (bob)',
             '403 — не организатор'],
        ],
        col_widths_cm=[3.0, 0.8, 6.5, 6.2],
    )

    # ╔══════════════════════════════════════════╗
    # ║   ПРИЛОЖЕНИЕ Б — Полный API-контракт     ║
    # ╚══════════════════════════════════════════╝
    add_heading_chapter(doc, 'ПРИЛОЖЕНИЕ Б')
    add_heading_centered(doc, 'Полный API-контракт')
    add_table(
        doc,
        ['№', 'URL', 'Метод', 'Назначение', 'Auth', 'Коды'],
        [
            ['1', '/api/v1/users/', 'POST', 'Регистрация (Djoser)', 'public',
             '201/400'],
            ['2', '/api/v1/users/me/', 'GET', 'Профиль текущего пользователя',
             'JWT', '200/401'],
            ['3', '/api/v1/auth/jwt/create/', 'POST', 'Получить access+refresh',
             'public', '200/401'],
            ['4', '/api/v1/auth/jwt/refresh/', 'POST', 'Обновить access',
             'public', '200/401'],
            ['5', '/api/v1/cats/', 'GET', 'Список котов с фильтрами и пагинацией',
             'public', '200'],
            ['6', '/api/v1/cats/', 'POST', 'Создать кота (owner = request.user)',
             'JWT', '201/400/401'],
            ['7', '/api/v1/cats/{id}/', 'GET', 'Карточка кота',
             'public', '200/404'],
            ['8', '/api/v1/cats/{id}/', 'PATCH / DELETE', 'Изменить / удалить '
             '(только владелец)', 'JWT', '200/204/401/403/404'],
            ['9', '/api/v1/cats/{id}/upload_image/', 'POST',
             'Загрузить фото (multipart)', 'JWT, владелец', '200/401/403'],
            ['10', '/api/v1/achievements/', 'GET / POST',
             'Список / создание (POST — staff)', 'public/staff', '200/201'],
            ['11', '/api/v1/tags/', 'GET', 'Список тегов', 'public', '200'],
            ['12', '/api/v1/events/', 'GET', 'Список ивентов',
             'public', '200'],
            ['13', '/api/v1/events/', 'POST', 'Создать ивент',
             'JWT', '201/400/401'],
            ['14', '/api/v1/events/{id}/', 'GET', 'Карточка ивента',
             'public', '200/404'],
            ['15', '/api/v1/events/{id}/', 'PATCH / DELETE',
             'Изменить / удалить (организатор)', 'JWT', '200/204/401/403/404'],
            ['16', '/api/v1/events/{id}/applications/', 'GET',
             'Заявки на ивент', 'JWT', '200/401/404'],
            ['17', '/api/v1/applications/', 'POST',
             'Подать заявку (signal post_save)', 'JWT', '201/400/401'],
            ['18', '/api/v1/applications/', 'GET',
             'Мои заявки', 'JWT', '200/401'],
            ['19', '/api/v1/applications/{id}/set_status/', 'POST',
             'Сменить статус', 'JWT, role-aware', '200/400/403/404'],
            ['20', '/api/v1/dialogs/{id}/messages/', 'GET / POST',
             'Чат по заявке', 'JWT, участник', '200/201/401/403/404'],
            ['21', '/api/v1/messages/{id}/mark_read/', 'POST',
             'Отметить прочитанным', 'JWT, не автор', '200/400/403'],
            ['22', '/api/schema/', 'GET', 'OpenAPI 3.0 JSON',
             'public', '200'],
            ['23', '/api/schema/swagger-ui/', 'GET', 'Swagger UI',
             'public', '200'],
            ['24', '/api/schema/redoc/', 'GET', 'ReDoc',
             'public', '200'],
        ],
        col_widths_cm=[1.0, 4.5, 2.6, 4.4, 2.3, 1.6],
    )

    # ╔══════════════════════════════════════════╗
    # ║   ПРИЛОЖЕНИЕ В — Конфигурации Docker     ║
    # ╚══════════════════════════════════════════╝
    add_heading_chapter(doc, 'ПРИЛОЖЕНИЕ В')
    add_heading_centered(doc, 'Конфигурации Docker')

    add_heading_section(doc, 'docker-compose.yml')
    add_code(
        doc,
        'services:\n'
        '  db:\n'
        '    image: postgres:16-alpine\n'
        '    env_file: .env\n'
        '    volumes:\n'
        '      - pg_data:/var/lib/postgresql/data\n'
        '    healthcheck:\n'
        '      test: ["CMD-SHELL", "pg_isready -U $${POSTGRES_USER}"]\n'
        '      interval: 5s\n'
        '      retries: 5\n'
        '\n'
        '  backend:\n'
        '    build: .\n'
        '    env_file: .env\n'
        '    depends_on:\n'
        '      db: { condition: service_healthy }\n'
        '    ports:\n'
        '      - "8000:8000"\n'
        '    volumes:\n'
        '      - static_volume:/app/static\n'
        '      - media_volume:/app/media\n'
        '\n'
        '  nginx:\n'
        '    image: nginx:1.25-alpine\n'
        '    profiles: ["with-nginx"]\n'
        '    ports:\n'
        '      - "80:80"\n'
        '    depends_on: [backend]\n'
        '    volumes:\n'
        '      - ./nginx/default.conf:/etc/nginx/conf.d/default.conf:ro\n'
        '      - static_volume:/usr/share/nginx/html/static:ro\n'
        '      - media_volume:/usr/share/nginx/html/media:ro\n'
        '\n'
        'volumes:\n'
        '  pg_data: {}\n'
        '  static_volume: {}\n'
        '  media_volume: {}\n'
    )

    add_heading_section(doc, 'nginx/default.conf')
    add_code(
        doc,
        'server {\n'
        '    listen 80;\n'
        '    server_name _;\n'
        '\n'
        '    location /static/ { alias /usr/share/nginx/html/static/; }\n'
        '    location /media/  { alias /usr/share/nginx/html/media/;  }\n'
        '\n'
        '    location / {\n'
        '        proxy_pass http://backend:8000;\n'
        '        proxy_set_header Host $host;\n'
        '        proxy_set_header X-Real-IP $remote_addr;\n'
        '        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;\n'
        '        proxy_set_header X-Forwarded-Proto $scheme;\n'
        '    }\n'
        '}\n'
    )

    # ───── СОХРАНЕНИЕ ─────
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    print(f'[ok] {OUT.name} собран ({OUT.stat().st_size:,} bytes)')


if __name__ == '__main__':
    build()
