"""Генерация отчёта по творческому заданию Kittygram (docx).

Структура — по шаблону «Творческое задание Kittygram.docx».
Скриншоты берутся из docs/screenshots/creative/.

Запуск:
    .venv/Scripts/python scripts/build_creative_report.py
Результат: ../../Any/PREU/.../Творческое_задание.docx
"""
import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


REPO = Path(__file__).resolve().parent.parent
SCREENS = REPO / 'docs' / 'screenshots' / 'creative'
OUT = (
    REPO.parent.parent.parent
    / 'Any' / 'PREU'
    / 'интеграция и управление приложениями на удаленном сервере'
    / 'Творческое_задание.docx'
)


def add_para(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, size=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    pf = p.paragraph_format
    pf.first_line_indent = Cm(1.25) if indent else Cm(0)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p.alignment = align
    return p


def add_heading_centered(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    return p


def add_heading_left(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.5
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0.5)
    pf.line_spacing = 1.15
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    return p


def style_table_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=12):
    cell.text = ''
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    p.alignment = align
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.15
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        style_table_cell(table.rows[0].cells[j], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            style_table_cell(table.rows[i].cells[j], str(val))
    if col_widths_cm:
        for row in table.rows:
            for j, w in enumerate(col_widths_cm):
                row.cells[j].width = Cm(w)
    return table


def add_image(doc, path, width_cm=14.0, caption=None):
    if not path.exists():
        add_para(doc, f'[ОТСУТСТВУЕТ ФАЙЛ СКРИНА: {path.name}]', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.first_line_indent = Cm(0)
        c.paragraph_format.line_spacing = 1.15
        cr = c.add_run(caption)
        cr.italic = True
        cr.font.name = 'Times New Roman'
        cr.font.size = Pt(12)


def setup_page(doc):
    section = doc.sections[0]
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')

    pPr = style.element.get_or_add_pPr()
    spc = pPr.find(qn('w:spacing'))
    if spc is None:
        from docx.oxml import OxmlElement
        spc = OxmlElement('w:spacing')
        pPr.append(spc)
    spc.set(qn('w:line'), '360')
    spc.set(qn('w:lineRule'), 'auto')


def build():
    doc = Document()
    setup_page(doc)

    # ───── Шапка ─────
    add_heading_centered(doc, 'ТВОРЧЕСКОЕ ЗАДАНИЕ ПО ПРОЕКТУ KITTYGRAM')
    add_heading_centered(doc, 'Расширение функциональности через REST API (Django REST Framework)')

    add_para(doc, 'Дисциплина: Интеграция и управление приложениями на удалённом сервере.', indent=False)
    add_para(doc, 'ФИО студента: Сергеев А. А.', indent=False)
    add_para(doc, 'Группа: ПИ2у/24б.', indent=False)
    add_para(doc, 'Преподаватель: Брызгалов А. А.', indent=False)
    add_para(
        doc,
        'Тема: Проектирование и реализация серверной части проекта Kittygram '
        'для поддержки пользовательского сценария «Мини-чат по заявкам».',
        indent=False,
    )
    add_para(
        doc,
        'Ссылка на репозиторий (URL): https://github.com/yoursAnthony/kitty '
        '(ветка feature/creative).',
        indent=False,
    )

    # ───── 1. Выбор темы ─────
    add_heading_left(doc, '1. Выбор темы')
    add_para(
        doc,
        'Из перечня предложенных треков выбран сценарий «Мини-чат по заявкам/ивентам». '
        'В качестве предметной области для заявок выбраны кото-ивенты — события, '
        'на которые владельцы котов могут подавать заявки на участие со своими питомцами '
        '(фотосессии, прогулки в парке, встречи владельцев британских котов, выставки). '
        'Каждая поданная заявка автоматически открывает диалог между организатором ивента '
        'и заявителем, где они уточняют детали участия.',
    )

    # ───── 2. Цель и краткое описание идеи ─────
    add_heading_left(doc, '2. Цель и краткое описание идеи')
    add_para(
        doc,
        'Цель — расширить базовый Kittygram пользовательским сценарием «организатор '
        '— участник ивента» с переписокой по заявке, реализованным средствами REST API '
        'на Django REST Framework.',
    )
    add_para(
        doc,
        'Сценарий использования. Авторизованный пользователь — организатор — создаёт '
        'кото-ивент с указанием места, времени, ёмкости (capacity). Другие пользователи '
        'видят ивент в публичном списке и подают заявку на участие, выбирая, с каким '
        'из своих котов они хотят прийти. Сразу после подачи заявки на стороне сервера '
        'через сигнал post_save автоматически создаётся диалог по этой заявке и в нём '
        'размещается первое сообщение от заявителя — это исходный текст обращения. '
        'Дальше организатор и заявитель ведут переписку по REST-эндпоинту, организатор '
        'может одобрить или отклонить заявку через кастомное действие, а заявитель — '
        'отозвать. При финализации заявки диалог автоматически закрывается, и попытка '
        'писать в него возвращает 403.',
    )

    # ───── 3. Обязательные требования к реализации (минимум) ─────
    add_heading_left(doc, '3. Обязательные требования к реализации')
    add_para(doc, 'Все обязательные пункты задания выполнены и подкреплены прогоном Postman-коллекции.')
    add_table(
        doc,
        ['Требование', 'Реализовано', 'Где'],
        [
            ['≥ 2 новые модели и связи',
             '4 модели: Event, Application, Dialog, Message',
             'events/models.py'],
            ['≥ 4 новых эндпоинта (помимо базового Kittygram)',
             '9 новых эндпоинтов (events, applications, dialogs, messages)',
             'events/urls.py'],
            ['≥ 1 кастомное действие',
             '2 action: POST /applications/{id}/set_status/, POST /messages/{id}/mark_read/',
             'events/views.py'],
            ['Permissions',
             'IsOrganizerOrReadOnly, IsApplicationParticipant, IsDialogParticipant',
             'events/permissions.py'],
            ['≥ 2 валидации',
             '7 правил: свой ивент, чужой кот, прошедшая дата, дубль (event,cat), '
             'capacity при approve, role-aware смена статуса, закрытый диалог',
             'events/serializers.py'],
            ['Фильтрация и пагинация в одном списочном эндпоинте',
             'GET /events/ + /applications/ — фильтры status, organizer, '
             'starts_after; пагинация 10/страница',
             'events/filters.py'],
            ['Документация API',
             'Swagger UI + ReDoc через drf-spectacular',
             '/api/schema/swagger-ui/, /api/schema/redoc/'],
            ['Секреты через переменные окружения; .env.example',
             '.env.example в корне репозитория',
             '.env.example'],
        ],
        col_widths_cm=[5, 7, 4.5],
    )

    # ───── 4. Проектирование данных ─────
    add_heading_left(doc, '4. Проектирование данных')

    add_heading_left(doc, '4.1. Таблица моделей')
    add_table(
        doc,
        ['Модель', 'Назначение', 'Ключевые поля', 'Связи', 'Ограничения'],
        [
            ['Event',
             'Кото-ивент (фотосессия, прогулка, выставка) с лимитом мест',
             'title, description, location, starts_at, ends_at, capacity',
             'FK organizer→User',
             'title 4..120; capacity 1..100; CHECK ends_at > starts_at; на create starts_at в будущем'],
            ['Application',
             'Заявка пользователя на участие в ивенте со своим котом',
             'message_text, status (pending/approved/rejected/cancelled)',
             'FK event→Event, FK cat→Cat, FK applicant→User',
             'UNIQUE(event, cat) — один кот не подаётся дважды'],
            ['Dialog',
             'Диалог по заявке (создаётся автоматически сигналом post_save)',
             'is_closed, created_at',
             'OneToOne application→Application',
             'Уникальность через OneToOne'],
            ['Message',
             'Сообщение в диалоге',
             'text, is_read, created_at',
             'FK dialog→Dialog, FK author→User',
             'text не пустой; ordering по created_at'],
        ],
        col_widths_cm=[2.5, 3, 3.5, 3, 4.5],
    )

    add_heading_left(doc, '4.2. Бизнес-правила и валидации')
    add_table(
        doc,
        ['Правило', 'Где проверяется', 'Что возвращаем при нарушении', 'Пример'],
        [
            ['Нельзя подавать заявку на собственный ивент',
             'ApplicationSerializer.validate (events/serializers.py)',
             '400 Bad Request: {"event": ["Вы организатор этого ивента — заявку подавать не нужно."]}',
             'alice (organizer) → POST /applications/ {event: own, ...}'],
            ['Нельзя подавать заявку с чужим котом',
             'ApplicationSerializer.validate',
             '400 Bad Request: {"cat": ["Можно подать заявку только со своим котом."]}',
             'bob → POST /applications/ {cat: cat_id_alice, ...}'],
            ['Нельзя подавать заявку на прошедший/начавшийся ивент',
             'ApplicationSerializer.validate',
             '400 Bad Request: {"event": ["Ивент уже начался — заявку подавать поздно."]}',
             'event.starts_at < now()'],
            ['Один кот = одна заявка на ивент (UNIQUE)',
             'Application.Meta.constraints + serializer pre-check',
             '400 Bad Request: {"cat": ["Заявка с этим котом на этот ивент уже подана."]}',
             'bob дублирует заявку с тем же котом'],
            ['Capacity не превышается при approve',
             'ApplicationStatusSerializer.validate_status',
             '400 Bad Request: ["На ивенте больше нет свободных мест."]',
             'organizer пробует одобрить когда уже capacity заявок approved'],
            ['Сменить статус: approve/reject — организатор; cancel — заявитель',
             'ApplicationStatusSerializer (events/serializers.py)',
             '400 Bad Request с пояснением, ИЛИ 403/404 при попытке доступа',
             'bob пытается approve чужую заявку'],
            ['Писать в закрытый диалог нельзя',
             'DialogMessagesViewSet.create',
             '403 Forbidden: {"detail": "Диалог закрыт — писать в него нельзя."}',
             'после approve диалог.is_closed=True, bob → POST /messages/'],
        ],
        col_widths_cm=[3, 3.5, 4.5, 4.5],
    )

    # ───── 5. Permissions ─────
    add_heading_left(doc, '5. Права доступа (permissions)')
    add_table(
        doc,
        ['Сущность/эндпоинт', 'Чтение', 'Создание/изменение', 'Кто может'],
        [
            ['/events/',
             'public',
             'JWT обязателен; PATCH/DELETE только организатор',
             'IsOrganizerOrReadOnly (events/permissions.py)'],
            ['/applications/',
             'JWT; видит только свои заявки + заявки на свои ивенты',
             'JWT обязателен',
             'IsAuthenticated; queryset фильтрует по applicant ИЛИ event.organizer'],
            ['/applications/{id}/set_status/',
             '—',
             'approve/reject → организатор; cancel → заявитель',
             'IsAuthenticated + IsApplicationParticipant + role-aware валидация'],
            ['/dialogs/{id}/messages/',
             'JWT, только участник',
             'JWT, участник, диалог не закрыт',
             'IsAuthenticated; DialogMessagesViewSet._get_dialog проверяет has_participant'],
            ['/messages/{id}/mark_read/',
             '—',
             'участник диалога, не автор сообщения',
             'IsAuthenticated + IsDialogParticipant + проверка author≠user'],
        ],
        col_widths_cm=[4, 3, 4, 5],
    )

    # ───── 6. API-контракт ─────
    add_heading_left(doc, '6. API-контракт (таблица эндпоинтов)')
    add_table(
        doc,
        ['№', 'URL', 'Метод', 'Описание', 'Auth/права', 'Коды'],
        [
            ['1', '/api/v1/events/', 'GET',
             'Список кото-ивентов с фильтром, поиском, пагинацией', 'public', '200'],
            ['2', '/api/v1/events/', 'POST',
             'Создать ивент (organizer = request.user)', 'JWT', '201 / 400 / 401'],
            ['3', '/api/v1/events/{id}/', 'GET / PATCH / DELETE',
             'CRUD ивента (PATCH/DELETE — только организатор)', 'JWT для write',
             '200 / 401 / 403 / 404'],
            ['4', '/api/v1/events/{id}/applications/', 'GET',
             'Заявки ивента (организатор видит все, остальные — только свои)',
             'JWT', '200 / 401 / 404'],
            ['5', '/api/v1/applications/', 'POST',
             'Подать заявку. Сигнал post_save создаёт Dialog + первое Message',
             'JWT', '201 / 400 / 401'],
            ['6', '/api/v1/applications/', 'GET',
             '«Мои» заявки (как заявитель + как организатор), фильтр по статусу',
             'JWT', '200 / 401'],
            ['7', '/api/v1/applications/{id}/set_status/', 'POST',
             'КАСТОМНЫЙ action — смена статуса (approve/reject/cancel)',
             'JWT, role-aware', '200 / 400 / 401 / 403 / 404'],
            ['8', '/api/v1/dialogs/{id}/messages/', 'GET / POST',
             'Сообщения диалога: чтение и отправка',
             'JWT, участник', '200 / 201 / 401 / 403 / 404'],
            ['9', '/api/v1/messages/{id}/mark_read/', 'POST',
             'КАСТОМНЫЙ action — отметить чужое сообщение прочитанным',
             'JWT, участник, не автор', '200 / 400 / 401 / 403 / 404'],
        ],
        col_widths_cm=[1, 4.2, 2, 4, 2.5, 1.8],
    )

    # ───── 6.1 JSON-примеры ─────
    add_heading_left(doc, '6.1. Примеры запросов и ответов')

    add_heading_left(doc, 'Создание ивента (успех)')
    add_code(
        doc,
        'POST /api/v1/events/ HTTP/1.1\n'
        'Authorization: Bearer <alice_access>\n'
        'Content-Type: application/json\n'
        '\n'
        '{\n'
        '  "title": "Фотосессия в парке Горького",\n'
        '  "description": "Будем фотографировать котов в осенней листве.",\n'
        '  "location": "Москва, Парк Горького, у Главного входа",\n'
        '  "starts_at": "2026-10-18T16:00:00Z",\n'
        '  "ends_at":   "2026-10-18T19:00:00Z",\n'
        '  "capacity": 5\n'
        '}',
    )
    add_para(doc, 'Ответ — HTTP 201 Created:', indent=False)
    add_code(
        doc,
        '{\n'
        '  "id": 1,\n'
        '  "organizer": "alice",\n'
        '  "title": "Фотосессия в парке Горького",\n'
        '  "description": "Будем фотографировать котов в осенней листве.",\n'
        '  "location": "Москва, Парк Горького, у Главного входа",\n'
        '  "starts_at": "2026-10-18T19:00:00+03:00",\n'
        '  "ends_at":   "2026-10-18T22:00:00+03:00",\n'
        '  "capacity": 5,\n'
        '  "approved_count": 0,\n'
        '  "created_at": "2026-05-03T16:34:31.678443+03:00",\n'
        '  "updated_at": "2026-05-03T16:34:31.678449+03:00"\n'
        '}',
    )

    add_heading_left(doc, 'Подача заявки + автоматическое создание диалога')
    add_code(
        doc,
        'POST /api/v1/applications/ HTTP/1.1\n'
        'Authorization: Bearer <bob_access>\n'
        'Content-Type: application/json\n'
        '\n'
        '{\n'
        '  "event": 1,\n'
        '  "cat": 3,\n'
        '  "message_text": "Привет! Хочу прийти со своим Мурзиком."\n'
        '}',
    )
    add_para(doc, 'Ответ — HTTP 201 Created (виден сразу dialog_id, созданный сигналом):', indent=False)
    add_code(
        doc,
        '{\n'
        '  "id": 1,\n'
        '  "event": 1, "event_title": "Фотосессия в парке Горького",\n'
        '  "cat": 3, "cat_name": "Мурзик",\n'
        '  "applicant": "bob",\n'
        '  "message_text": "Привет! Хочу прийти со своим Мурзиком.",\n'
        '  "status": "pending",\n'
        '  "dialog_id": 1,\n'
        '  "created_at": "2026-05-03T16:34:31.889827+03:00",\n'
        '  "updated_at": "2026-05-03T16:34:31.889841+03:00"\n'
        '}',
    )

    add_heading_left(doc, 'Кастомный action set_status (одобрение организатором)')
    add_code(
        doc,
        'POST /api/v1/applications/1/set_status/ HTTP/1.1\n'
        'Authorization: Bearer <alice_access>\n'
        '\n'
        '{ "status": "approved" }\n'
        '\n'
        '→ HTTP 200 OK; status=approved; диалог автоматически закрывается.',
    )

    add_heading_left(doc, 'Ошибка: чужой кот (400)')
    add_code(
        doc,
        'POST /api/v1/applications/  body={"event":1,"cat":1,"message_text":"..."}\n'
        '(cat id=1 принадлежит alice, заявку подаёт bob)\n'
        '\n'
        'HTTP 400 Bad Request\n'
        '{ "cat": ["Можно подать заявку только со своим котом."] }',
    )

    add_heading_left(doc, 'Ошибка: попытка писать в закрытый диалог (403)')
    add_code(
        doc,
        'POST /api/v1/dialogs/1/messages/  body={"text":"Спасибо!"}\n'
        '(заявка одобрена, диалог закрыт)\n'
        '\n'
        'HTTP 403 Forbidden\n'
        '{ "detail": "Диалог закрыт — писать в него нельзя." }',
    )

    # ───── Скрины Swagger / ReDoc ─────
    add_image(doc, SCREENS / '01_swagger_top.png',
              caption='Рисунок 1 — Swagger UI: заголовок API и начало списка эндпоинтов (видны /applications/ и кастомный set_status)')
    add_image(doc, SCREENS / '02_swagger_dialogs_events.png',
              caption='Рисунок 2 — Swagger UI: эндпоинты /dialogs/{id}/messages/, /events/, /events/{id}/applications/, /messages/{id}/mark_read/')
    add_image(doc, SCREENS / '03_swagger_application_post.png',
              caption='Рисунок 3 — Swagger UI: раскрытый POST /applications/ — схема тела (event, cat, message_text) и пример ответа 201 с автосозданным dialog_id')
    add_image(doc, SCREENS / '04_redoc_events_create.png',
              caption='Рисунок 4 — ReDoc: схема POST /events/ — title, description, location, starts_at, ends_at, capacity (с ограничениями)')
    add_image(doc, SCREENS / '05_redoc_set_status.png',
              caption='Рисунок 5 — ReDoc: кастомное действие POST /applications/{id}/set_status/ — request body и пример 200-ответа')

    # ───── 7. Фильтрация и пагинация ─────
    add_heading_left(doc, '7. Фильтрация и пагинация')
    add_para(
        doc,
        'Реализованы фильтры на двух списочных эндпоинтах. Используется пакет '
        'django-filter, кроме него подключены SearchFilter и OrderingFilter из DRF. '
        'Пагинация — PageNumberPagination, по умолчанию 10 элементов на страницу.',
    )
    add_table(
        doc,
        ['Эндпоинт', 'Фильтры', 'Поиск/сортировка', 'Пример запроса'],
        [
            ['GET /api/v1/events/',
             '?organizer=, ?starts_after=, ?starts_before=',
             'search по title/location; ordering по starts_at, created_at, capacity',
             '/api/v1/events/?organizer=alice&starts_after=2026-09-01T00:00:00Z&ordering=-starts_at'],
            ['GET /api/v1/applications/',
             '?status=, ?event=',
             'ordering по created_at, status',
             '/api/v1/applications/?status=approved&ordering=-created_at'],
        ],
        col_widths_cm=[3.5, 3.5, 3.5, 5.5],
    )

    # ───── 8. Проверка через Postman ─────
    add_heading_left(doc, '8. Проверка через Postman / newman')
    add_para(
        doc,
        'Подготовлена коллекция postman/kittygram_creative.postman_collection.json '
        '(формат v2.1) — 17 запросов, покрывающих весь сценарий: получение JWT '
        'для alice/bob/charlie, создание ивента, поиск, подача заявки, негативные '
        'кейсы (свой ивент, чужой кот, дубль), переписка в диалоге, попытка постороннего '
        'читать чужой диалог, кастомные actions set_status и mark_read, попытка '
        'писать в закрытый диалог, фильтр applications по статусу, nested список '
        'заявок ивента.',
    )
    add_para(
        doc,
        'Прогон через newman v6: 18 запросов, 16 test-scripts, 22 ассерта — все зелёные, '
        '0 failed.',
    )
    add_code(
        doc,
        '$ newman run postman/kittygram_creative.postman_collection.json \\\n'
        '       --env-var "base_url=http://127.0.0.1:8000"\n'
        '...\n'
        '│              iterations │   1 │   0 │\n'
        '│                requests │  18 │   0 │\n'
        '│            test-scripts │  16 │   0 │\n'
        '│              assertions │  22 │   0 │\n'
        '│ total run duration: 2.8s ...                                     │',
    )
    add_image(doc, SCREENS / '10_newman_top.png', width_cm=15,
              caption='Рисунок 6 — Newman HTML-репорт: сводка прогона (18 запросов, 22 ассерта, 0 failed)')
    add_image(doc, SCREENS / '11_newman_middle.png', width_cm=15,
              caption='Рисунок 7 — Newman: запросы 11–14 — charlie тестит чужой диалог (404), bob отмечает прочитанным (mark_read), alice одобряет (set_status)')
    add_image(doc, SCREENS / '12_newman_bottom.png', width_cm=15,
              caption='Рисунок 8 — Newman: запросы 14–16 — set_status (approved) + 403 при попытке писать в закрытый диалог + фильтр applications?status=approved')

    # ───── Админка ─────
    add_heading_left(doc, '9. Демонстрация в Django-админке')
    add_para(
        doc,
        'Все четыре модели расширения зарегистрированы в админке (events/admin.py) '
        'с list_display, list_filter, search_fields и autocomplete_fields для удобной '
        'модерации. На скринах ниже видны живые данные из PostgreSQL под docker-compose, '
        'засеянные scripts/seed.py.',
    )
    add_image(doc, SCREENS / '06_admin_index.png',
              caption='Рисунок 9 — Главная страница админки: блок «События и заявки» с разделами Диалоги / Заявки на ивенты / Кото-ивенты / Сообщения')
    add_image(doc, SCREENS / '07_admin_events.png',
              caption='Рисунок 10 — Список кото-ивентов: «Фотосессия в парке Горького» (alice) и «Встреча владельцев британских котов» (bob)')
    add_image(doc, SCREENS / '08_admin_applications.png',
              caption='Рисунок 11 — Заявки на ивенты: видны статусы (одобрена / на рассмотрении), коты (Барсик, Мурзик), заявители (alice, bob), фильтр справа по статусу')
    add_image(doc, SCREENS / '09_admin_messages.png',
              caption='Рисунок 12 — Сообщения в диалогах: 4 сообщения — первое автосообщение (от сигнала) и переписка организатора с заявителем')

    # ───── 10. Самооценка ─────
    add_heading_left(doc, '10. Самооценка по критериям')
    add_table(
        doc,
        ['Критерий', 'Макс.', 'Что сделано в проекте'],
        [
            ['Оригинальность идеи',
             '6',
             'Кото-ивенты + заявки + автодиалог через сигнал — целостная связка '
             '«организатор↔участник» с естественной переписокой и закрытием диалога '
             'при финализации заявки. Не «копия» простого CRUD.'],
            ['Цельность и полезность',
             '5',
             'Полный пользовательский путь: создание ивента → поиск ивентов → подача '
             'заявки → переписка → одобрение → закрытие диалога. Работает как единая '
             'система, а не набор таблиц.'],
            ['Дизайн API и качество реализации',
             '5',
             'REST-семантика, корректные HTTP-статусы (200/201/400/401/403/404), '
             '7 валидаций, 2 кастомных action, role-aware permissions, auto-сигнал, '
             'фильтрация и пагинация на 2 эндпоинтах.'],
            ['Презентация и полировка',
             '4',
             'Swagger UI + ReDoc на drf-spectacular; Postman-коллекция (17 запросов, '
             '22 ассерта зелёные); voспроизводимость через docker compose up; .env.example; '
             'демо-данные через seed.py.'],
        ],
        col_widths_cm=[5, 1.5, 9.5],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    sys.stdout.reconfigure(encoding='utf-8')
    print(f'OK: {OUT} ({OUT.stat().st_size} bytes)')


if __name__ == '__main__':
    build()
