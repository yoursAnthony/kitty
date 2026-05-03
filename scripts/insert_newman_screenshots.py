"""Edit-in-place: добавляет в Курсовую_работу.docx два недостающих скриншота
newman (запросы 4–7 и 8–10) между текущими Рис. 12 и Рис. 13, обновляет
нумерацию последующих рисунков (13→15, 14→16, 15→17) и текстовые упоминания.

ВАЖНО: не перегенерирует документ из шаблона — сохраняет ручные правки
пользователя в титульнике.

Запуск (из корня проекта):
    .venv/Scripts/python scripts/insert_newman_screenshots.py
"""
import io
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


REPO = Path(__file__).resolve().parent.parent
SCREENS = REPO / 'docs' / 'screenshots' / 'creative'
DOCX = (
    REPO.parent.parent.parent
    / 'Any' / 'PREU'
    / 'интеграция и управление приложениями на удаленном сервере'
    / 'Курсовая_работа.docx'
)

# Какие подписи переименовать (старая → новая)
RENUMBER = [
    ('Рисунок 15 — Deployment',   'Рисунок 17 — Deployment'),
    ('Рисунок 14 — Newman',       'Рисунок 16 — Newman'),
    ('Рисунок 13 — Newman',       'Рисунок 15 — Newman'),
]

# Какие фразы в основном тексте обновить
TEXT_REPLACEMENTS = [
    ('На рисунках 12–14 показаны фрагменты HTML-отчёта newman',
     'На рисунках 12–16 показаны фрагменты HTML-отчёта newman'),
    ('Полная схема развёртывания приложения приведена на рисунке 15',
     'Полная схема развёртывания приложения приведена на рисунке 17'),
]

# Куда вставлять (после параграфа подписи Рис. 12 newman_top)
ANCHOR_CAPTION_PREFIX = 'Рисунок 12 — Newman'

# Что вставить (img_path, caption)
NEW_FIGURES = [
    (SCREENS / '10b_newman_4_7.png',
     'Рисунок 13 — Newman: фрагмент отчёта с запросами 4–7'),
    (SCREENS / '10c_newman_8_10.png',
     'Рисунок 14 — Newman: фрагмент отчёта с запросами 8–10'),
]


def find_caption_paragraph(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def replace_text_in_runs(paragraph, old, new):
    """Заменяет подстроку, итерируясь по run’ам (сохраняет форматирование).
    Работает корректно, если old целиком умещается в одном run’е,
    что выполняется для подписей рисунков (один run на параграф).
    """
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # Fallback: соединить все runs и распределить обратно одним
    full = ''.join(r.text for r in paragraph.runs)
    if old in full:
        new_full = full.replace(old, new)
        # Стереть все runs, оставив один с заменённым текстом
        first = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        first.text = new_full
        for r in paragraph.runs[1:]:
            r.text = ''
        return True
    return False


def add_image_paragraph(doc, img_path, width_cm=15.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))
    return p


def add_caption_paragraph(doc, caption_text):
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.first_line_indent = Cm(0)
    c.paragraph_format.line_spacing = 1.15
    c.paragraph_format.space_after = Pt(6)
    cr = c.add_run(caption_text)
    cr.italic = True
    cr.font.name = 'Times New Roman'
    cr.font.size = Pt(12)
    return c


def main():
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

    if not DOCX.exists():
        print(f'[err] Не найден {DOCX}')
        sys.exit(1)
    for img_path, _ in NEW_FIGURES:
        if not img_path.exists():
            print(f'[err] Нет файла скрина: {img_path}')
            sys.exit(1)

    doc = Document(str(DOCX))

    # 1. Сначала переименуем последующие рисунки 13/14/15 → 15/16/17.
    #    Делаем в обратном порядке (15→17, 14→16, 13→15), чтобы новые названия
    #    не пересекались со старыми префиксами.
    for old_prefix, new_prefix in RENUMBER:
        for p in doc.paragraphs:
            if old_prefix in p.text:
                if replace_text_in_runs(p, old_prefix, new_prefix):
                    print(f'[ok] переименовано: «{old_prefix}» → «{new_prefix}»')
                    break

    # 2. Заменим текстовые упоминания
    for old, new in TEXT_REPLACEMENTS:
        for p in doc.paragraphs:
            if old in p.text:
                if replace_text_in_runs(p, old, new):
                    print(f'[ok] текст: «{old[:50]}…» обновлён')
                    break

    # 3. Найдём anchor — параграф с подписью Рис. 12 newman_top
    anchor = find_caption_paragraph(doc, ANCHOR_CAPTION_PREFIX)
    if anchor is None:
        print(f'[err] Не найден anchor «{ANCHOR_CAPTION_PREFIX}»')
        sys.exit(1)
    print(f'[ok] anchor найден: «{anchor.text.strip()}»')

    # 4. Добавим в КОНЕЦ документа image + caption для каждого нового рисунка,
    #    затем перенесём блоки сразу после anchor (в обратном порядке).
    new_blocks = []  # список (image_paragraph, caption_paragraph)
    for img_path, caption in NEW_FIGURES:
        img_p = add_image_paragraph(doc, img_path, width_cm=15.0)
        cap_p = add_caption_paragraph(doc, caption)
        new_blocks.append((img_p, cap_p))

    # Перенос: anchor.addnext(блок2.cap), anchor.addnext(блок2.img),
    #         anchor.addnext(блок1.cap), anchor.addnext(блок1.img)
    # Так после anchor окажется: img1, cap1, img2, cap2
    for img_p, cap_p in reversed(new_blocks):
        anchor._element.addnext(cap_p._element)
        anchor._element.addnext(img_p._element)
    print(f'[ok] вставлено {len(NEW_FIGURES)} новых рисунков после anchor')

    doc.save(str(DOCX))
    print(f'[ok] сохранён: {DOCX.name} ({DOCX.stat().st_size:,} bytes)')


if __name__ == '__main__':
    main()
