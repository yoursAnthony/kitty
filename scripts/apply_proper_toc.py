"""Edit-in-place: заменяет статическое «Содержание» на настоящее TOC-поле Word
и добавляет сквозную нумерацию страниц (начиная со 2-й страницы — титульник без номера).

Что делает:
1. Перенастраивает стили Heading 1, Heading 2 под оформление РЭУ
   (Times New Roman 14, чёрный, без курсива; H1 — центр + page-break-before;
    H2 — слева).
2. Применяет стиль Heading 1 ко всем главным заголовкам:
   ВВЕДЕНИЕ, ГЛАВА N, ЗАКЛЮЧЕНИЕ, СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ,
   ПРИЛОЖЕНИЕ А|Б|В.
3. Применяет стиль Heading 2 ко всем подразделам формата «1.1.», «1.2.», ...
4. Удаляет старое «ручное» содержание (47 строк с точками).
5. Вставляет настоящий TOC field на место содержания.
6. Включает «different first page» в секции и кладёт в footer
   PAGE field — на титульнике номер не выводится, со 2-й страницы — да.

Запуск (из корня проекта):
    .venv/Scripts/python scripts/apply_proper_toc.py

После открытия в Word — нажмите Ctrl+A → F9, выберите «Обновить целиком».
"""
import io
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPO = Path(__file__).resolve().parent.parent
DOCX = (
    REPO.parent.parent.parent
    / 'Any' / 'PREU'
    / 'интеграция и управление приложениями на удаленном сервере'
    / 'Курсовая_работа.docx'
)


# Регексы определения заголовков
H1_PATTERNS = [
    r'^СОДЕРЖАНИЕ$',
    r'^ВВЕДЕНИЕ$',
    r'^ГЛАВА \d+\.\s.*$',
    r'^ЗАКЛЮЧЕНИЕ$',
    r'^СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ$',
    r'^ПРИЛОЖЕНИЕ [А-Я]$',
]
H2_PATTERN = r'^\d+\.\d+\.\s+\S'


def is_h1(text):
    return any(re.match(p, text) for p in H1_PATTERNS)


def is_h2(text):
    return bool(re.match(H2_PATTERN, text))


def configure_heading_styles(doc):
    """Настраивает стили Heading 1 и Heading 2 под оформление РЭУ."""
    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.italic = False
    h1.font.color.rgb = RGBColor(0, 0, 0)
    pf1 = h1.paragraph_format
    pf1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf1.line_spacing = 1.5
    pf1.first_line_indent = Cm(0)
    pf1.space_before = Pt(12)
    pf1.space_after = Pt(12)
    pf1.page_break_before = True
    pf1.keep_with_next = True

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.italic = False
    h2.font.color.rgb = RGBColor(0, 0, 0)
    pf2 = h2.paragraph_format
    pf2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf2.line_spacing = 1.5
    pf2.first_line_indent = Cm(0)
    pf2.space_before = Pt(8)
    pf2.space_after = Pt(4)
    pf2.page_break_before = False
    pf2.keep_with_next = True


def apply_heading_styles(doc):
    """Назначает Heading 1 / Heading 2 параграфам, чьи тексты соответствуют шаблонам."""
    h1_count = 0
    h2_count = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if is_h1(text):
            p.style = doc.styles['Heading 1']
            h1_count += 1
        elif is_h2(text):
            p.style = doc.styles['Heading 2']
            h2_count += 1
    print(f'[ok] Heading 1 применён к {h1_count} параграфам')
    print(f'[ok] Heading 2 применён к {h2_count} параграфам')


def find_paragraph_by_text(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None


def remove_paragraph(paragraph):
    el = paragraph._element
    el.getparent().remove(el)


def replace_static_toc_with_field(doc):
    """Удаляет старое статическое содержание и вставляет TOC field."""
    soderzhanie = find_paragraph_by_text(doc, 'СОДЕРЖАНИЕ')
    vvedenie = find_paragraph_by_text(doc, 'ВВЕДЕНИЕ')
    if soderzhanie is None or vvedenie is None:
        print('[err] Не найден «СОДЕРЖАНИЕ» или «ВВЕДЕНИЕ» — пропуск замены TOC.')
        return

    # Соберём элементы между soderzhanie и vvedenie через XML-обход (XML стабилен,
    # в отличие от paragraphs.index, который не находит объект из нового списка).
    sod_el = soderzhanie._element
    vved_el = vvedenie._element
    parent = sod_el.getparent()

    to_delete = []
    cur = sod_el.getnext()
    while cur is not None and cur is not vved_el:
        to_delete.append(cur)
        cur = cur.getnext()
    print(f'[ok] Удаляется {len(to_delete)} элементов старого содержания')
    for el in to_delete:
        parent.remove(el)

    # Создаём новый параграф с TOC field в конце документа,
    # затем перенесём его перед «ВВЕДЕНИЕ».
    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.first_line_indent = Cm(0)
    toc_p.paragraph_format.line_spacing = 1.5
    toc_p.paragraph_format.space_before = Pt(6)
    toc_p.paragraph_format.space_after = Pt(6)

    # XML TOC field: { TOC \o "1-2" \h \z \u }
    p_el = toc_p._element

    # Run 1: fldChar begin
    r1 = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '28')
    rPr.append(sz)
    r1.append(rPr)
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r1.append(fld_begin)
    p_el.append(r1)

    # Run 2: instrText
    r2 = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    r2.append(instr)
    p_el.append(r2)

    # Run 3: fldChar separate
    r3 = OxmlElement('w:r')
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    r3.append(fld_sep)
    p_el.append(r3)

    # Run 4: placeholder text (показывается до первого обновления F9)
    r4 = OxmlElement('w:r')
    rPr4 = OxmlElement('w:rPr')
    rFonts4 = OxmlElement('w:rFonts')
    rFonts4.set(qn('w:ascii'), 'Times New Roman')
    rFonts4.set(qn('w:hAnsi'), 'Times New Roman')
    rPr4.append(rFonts4)
    sz4 = OxmlElement('w:sz')
    sz4.set(qn('w:val'), '28')
    rPr4.append(sz4)
    italic = OxmlElement('w:i')
    rPr4.append(italic)
    r4.append(rPr4)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = 'Содержание будет создано при обновлении (Ctrl+A → F9 в Word).'
    r4.append(t)
    p_el.append(r4)

    # Run 5: fldChar end
    r5 = OxmlElement('w:r')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r5.append(fld_end)
    p_el.append(r5)

    # Перенос перед ВВЕДЕНИЕ
    vvedenie._element.addprevious(p_el)
    print('[ok] TOC field вставлен после «СОДЕРЖАНИЕ»')


def add_page_numbers(doc):
    """Добавляет нумерацию страниц в footer (центр, TNR 12).
    На первой странице (титульник) номер не отображается."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Очищаем first-page footer (на титульнике номера нет)
    fp_footer = section.first_page_footer
    for p in list(fp_footer.paragraphs):
        # Очищаем содержимое, оставляя пустой параграф
        for r in list(p.runs):
            r.text = ''

    # Основной footer — со 2-й страницы — содержит PAGE field
    footer = section.footer
    if footer.paragraphs:
        footer_p = footer.paragraphs[0]
        # Очистим существующий контент
        for r in list(footer_p.runs):
            r._element.getparent().remove(r._element)
    else:
        footer_p = footer.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.first_line_indent = Cm(0)

    # Вставляем PAGE field
    p_el = footer_p._element

    r1 = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')  # 12pt
    rPr.append(sz)
    r1.append(rPr)
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r1.append(fld_begin)
    p_el.append(r1)

    r2 = OxmlElement('w:r')
    rPr2 = OxmlElement('w:rPr')
    rFonts2 = OxmlElement('w:rFonts')
    rFonts2.set(qn('w:ascii'), 'Times New Roman')
    rFonts2.set(qn('w:hAnsi'), 'Times New Roman')
    rPr2.append(rFonts2)
    sz2 = OxmlElement('w:sz')
    sz2.set(qn('w:val'), '24')
    rPr2.append(sz2)
    r2.append(rPr2)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE \\* MERGEFORMAT'
    r2.append(instr)
    p_el.append(r2)

    r3 = OxmlElement('w:r')
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    r3.append(fld_sep)
    p_el.append(r3)

    r4 = OxmlElement('w:r')
    rPr4 = OxmlElement('w:rPr')
    rFonts4 = OxmlElement('w:rFonts')
    rFonts4.set(qn('w:ascii'), 'Times New Roman')
    rFonts4.set(qn('w:hAnsi'), 'Times New Roman')
    rPr4.append(rFonts4)
    sz4 = OxmlElement('w:sz')
    sz4.set(qn('w:val'), '24')
    rPr4.append(sz4)
    r4.append(rPr4)
    t = OxmlElement('w:t')
    t.text = '2'
    r4.append(t)
    p_el.append(r4)

    r5 = OxmlElement('w:r')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r5.append(fld_end)
    p_el.append(r5)

    print('[ok] Нумерация страниц добавлена (титульник — без номера)')


def main():
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

    if not DOCX.exists():
        print(f'[err] Не найден {DOCX}')
        sys.exit(1)

    doc = Document(str(DOCX))
    configure_heading_styles(doc)
    # Важен порядок: сначала удаляем старое статическое содержание, иначе
    # его строки формата «1.1. ... 5» сматчат шаблон Heading 2.
    replace_static_toc_with_field(doc)
    apply_heading_styles(doc)
    add_page_numbers(doc)

    doc.save(str(DOCX))
    print(f'[ok] сохранён: {DOCX.name} ({DOCX.stat().st_size:,} bytes)')
    print('[i] Откройте в Word, выделите весь документ (Ctrl+A) и нажмите F9 — '
          'выберите «Обновить целиком». Содержание заполнится автоматически.')


if __name__ == '__main__':
    main()
