"""Генерация отчёта по заданию на текущий рейтинг (docx).

Структура — по шаблону «Отчет по заданию на текущий рейтинг.docx».
Скриншоты берутся из docs/screenshots/.

Запуск:
    .venv/Scripts/python scripts/build_current_rating_report.py
Результат: ../../Any/PREU/.../Отчёт_текущий_рейтинг.docx
"""
import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPO = Path(__file__).resolve().parent.parent
SCREENS = REPO / 'docs' / 'screenshots'
OUT = (
    REPO.parent.parent.parent
    / 'Any' / 'PREU'
    / 'интеграция и управление приложениями на удаленном сервере'
    / 'Отчёт_текущий_рейтинг.docx'
)


def set_default_paragraph(p, font_name='Times New Roman', font_size=14, bold=False, align=None):
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    for run in p.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            from docx.oxml import OxmlElement
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:cs'), font_name)


def add_para(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, size=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    pf = p.paragraph_format
    pf.first_line_indent = Cm(1.25) if indent else Cm(0)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    p.alignment = align
    return p


def add_heading_centered(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.5
    pf.space_before = Pt(12 if level == 1 else 6)
    pf.space_after = Pt(6)
    return p


def add_heading_left(doc, text, level=2):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.5
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0.5)
    pf.line_spacing = 1.15
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    return p


def style_table_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=12):
    cell.text = ''
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    p.alignment = align
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.15
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        style_table_cell(table.rows[0].cells[j], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            style_table_cell(table.rows[i].cells[j], str(val))
    if col_widths_cm:
        for row in table.rows:
            for j, w in enumerate(col_widths_cm):
                row.cells[j].width = Cm(w)
    return table


def add_image(doc, path, width_cm=15.0, caption=None):
    if not path.exists():
        add_para(doc, f'[ОТСУТСТВУЕТ ФАЙЛ СКРИНА: {path.name}]', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.first_line_indent = Cm(0)
        c.paragraph_format.line_spacing = 1.15
        cr = c.add_run(caption)
        cr.italic = True
        cr.font.name = 'Times New Roman'
        cr.font.size = Pt(12)


def setup_page(doc):
    section = doc.sections[0]
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')

    pPr = style.element.get_or_add_pPr()
    spc = pPr.find(qn('w:spacing'))
    if spc is None:
        from docx.oxml import OxmlElement
        spc = OxmlElement('w:spacing')
        pPr.append(spc)
    spc.set(qn('w:line'), '360')  # 1.5 интервал = 360 twentieths
    spc.set(qn('w:lineRule'), 'auto')


# ───────────────────────────── Содержимое ─────────────────────────────

def build():
    doc = Document()
    setup_page(doc)

    # Заголовок
    add_heading_centered(doc, 'ОТЧЁТ ПО ВЫПОЛНЕНИЮ ПРАКТИКУМА', level=1)

    add_para(doc, 'Дисциплина: Интеграция и управление приложениями на удалённом сервере.', indent=False)
    add_para(doc, 'ФИО студента: Сергеев А. А.', indent=False)
    add_para(doc, 'Группа: ПИ2у/24б.', indent=False)
    add_para(doc, 'Преподаватель: Брызгалов А. А.', indent=False)
    add_para(doc, 'Ссылка на репозиторий (URL): https://github.com/yoursAnthony/kitty', indent=False)

    # 1
    add_heading_left(doc, '1. Цель и итоговый результат')
    add_para(
        doc,
        'Учебный проект Kittygram расширен до полноценного REST API на стеке '
        'Django 4.2 LTS и Django REST Framework 3.14: внедрена JWT-аутентификация '
        '(Djoser + simplejwt), permissions с разграничением доступа по владельцу, '
        'фильтрация и пагинация, документация API через drf-spectacular '
        '(Swagger UI и ReDoc), throttling. Проект упакован в Docker-образ и '
        'разворачивается через docker-compose в связке PostgreSQL 16 + Gunicorn '
        '(+ опциональный Nginx через профиль). Подготовлены README, .env.example '
        'и Postman-коллекция из 13 запросов, прогон которой завершается без ошибок.',
    )

    add_heading_left(doc, '1.1. Запуск через docker compose')
    add_code(
        doc,
        '$ docker compose up -d --build\n'
        ' Network kitty_default                Created\n'
        ' Volume "kitty_pg_data"               Created\n'
        ' Volume "kitty_static_volume"         Created\n'
        ' Volume "kitty_media_volume"          Created\n'
        ' Container kittygram_db               Started\n'
        ' Container kittygram_db               Healthy\n'
        ' Container kittygram_backend          Started\n'
        '\n'
        '$ docker compose ps\n'
        'NAME                IMAGE                STATUS                    PORTS\n'
        'kittygram_backend   kitty-backend        Up                        0.0.0.0:8000->8000/tcp\n'
        'kittygram_db        postgres:16-alpine   Up (healthy)              5432/tcp',
    )
    add_image(doc, SCREENS / '01_swagger_ui.png', width_cm=15,
              caption='Рисунок 1 — Swagger UI после запуска через docker-compose: заголовок Kittygram API и часть списка эндпоинтов')

    # 2
    add_heading_left(doc, '2. Модель данных и связи')

    add_heading_left(doc, '2.1. Таблица моделей')
    add_table(
        doc,
        ['Модель', 'Назначение', 'Ключевые поля', 'Связи', 'Ограничения'],
        [
            ['Cat',
             'Профиль кота',
             'name, color, birth_year, description, image, created_at',
             'FK owner→User; M2M tags→Tag (через CatTag); M2M achievements→Achievement (через CatAchievement)',
             'name 2..32; birth_year ∈ [1990; текущий год]; UNIQUE (owner, name)'],
            ['Tag',
             'Тег характера/настроения',
             'name, slug',
             'M2M cats→Cat',
             'name 2..64 unique; slug — regex ^[-a-zA-Z0-9_]+$, unique'],
            ['Achievement',
             'Справочник достижений',
             'name',
             'M2M cats→Cat',
             'name 1..64 unique'],
            ['CatTag',
             'Промежуточная таблица',
             'cat, tag',
             'FK cat, FK tag',
             'UNIQUE (cat, tag)'],
            ['CatAchievement',
             'Связь кота и достижения с датой получения',
             'cat, achievement, achieved_at',
             'FK cat, FK achievement',
             'UNIQUE (cat, achievement); achieved_at = auto_now_add'],
        ],
        col_widths_cm=[2.5, 3.5, 4, 3.5, 4.5],
    )

    add_heading_left(doc, '2.2. Бизнес-правила и валидации')
    add_table(
        doc,
        ['Правило', 'Где проверяется', 'Что возвращаем при нарушении', 'Пример'],
        [
            ['birth_year ∈ [1990; текущий год]',
             'model field validators (MinValueValidator/MaxValueValidator), Cat.birth_year',
             '400 Bad Request, тело: {"birth_year": ["Убедитесь, что это значение больше либо равно 1990."]}',
             'POST /api/v1/cats/ {"name":"BadYear","color":"black","birth_year":1800}'],
            ['Уникальность клички у одного владельца',
             'serializer.validate_name() в CatSerializer',
             '400 Bad Request, тело: {"name": ["У вас уже есть кот с такой кличкой."]}',
             'POST /api/v1/cats/ {"name":"Pushok",...} при уже существующем Pushok у этого пользователя'],
            ['Slug тега — только латиница/цифры/_/-',
             'RegexValidator на Tag.slug',
             '400 Bad Request с подсказкой о допустимых символах',
             'POST /api/v1/tags/ {"slug":"кот"} (некорректно)'],
            ['Анонимный пользователь не может изменять данные',
             'IsAuthenticatedOrReadOnly + IsOwnerOrReadOnly (DRF permissions)',
             '401 Unauthorized, тело: {"detail":"Учётные данные не были предоставлены."}',
             'POST /api/v1/cats/ без заголовка Authorization'],
        ],
        col_widths_cm=[3.5, 3.5, 4.5, 4.5],
    )

    # 3
    add_heading_left(doc, '3. API-контракт (эндпоинты)')
    add_table(
        doc,
        ['№', 'URL', 'Метод', 'Описание', 'Auth/права', 'Коды'],
        [
            ['1', '/api/v1/auth/users/', 'POST', 'Регистрация пользователя', 'public', '201 / 400'],
            ['2', '/api/v1/auth/users/me/', 'GET', 'Профиль текущего пользователя', 'JWT', '200 / 401'],
            ['3', '/api/v1/auth/jwt/create/', 'POST', 'Получить JWT (access + refresh)', 'public', '200 / 401'],
            ['4', '/api/v1/auth/jwt/refresh/', 'POST', 'Обновить access по refresh', 'public', '200 / 401'],
            ['5', '/api/v1/cats/', 'GET', 'Список котов с фильтром, поиском, пагинацией', 'public', '200'],
            ['6', '/api/v1/cats/', 'POST', 'Создать кота (owner = request.user)', 'JWT', '201 / 400 / 401'],
            ['7', '/api/v1/cats/{id}/', 'GET', 'Карточка кота', 'public', '200 / 404'],
            ['8', '/api/v1/cats/{id}/', 'PATCH', 'Частичное обновление', 'JWT, owner', '200 / 400 / 401 / 403 / 404'],
            ['9', '/api/v1/cats/{id}/', 'DELETE', 'Удалить кота', 'JWT, owner', '204 / 401 / 403 / 404'],
            ['10', '/api/v1/cats/{id}/upload_image/', 'POST', 'Кастомный action — загрузка фото (multipart)', 'JWT, owner', '200 / 400 / 401 / 403'],
            ['11', '/api/v1/achievements/', 'GET', 'Список достижений', 'public', '200'],
            ['12', '/api/v1/achievements/', 'POST', 'Создать достижение', 'JWT, staff', '201 / 401 / 403'],
            ['13', '/api/v1/tags/', 'GET', 'Список тегов', 'public', '200'],
            ['14', '/api/schema/swagger-ui/', 'GET', 'Swagger UI', 'public', '200'],
            ['15', '/api/schema/redoc/', 'GET', 'ReDoc', 'public', '200'],
        ],
        col_widths_cm=[1, 4.5, 1.5, 5, 2.5, 2],
    )

    add_heading_left(doc, '3.1. Примеры запросов и ответов')

    add_heading_left(doc, 'Успешное создание кота', level=3)
    add_code(
        doc,
        'POST /api/v1/cats/ HTTP/1.1\n'
        'Host: localhost:8000\n'
        'Authorization: Bearer <access_token>\n'
        'Content-Type: application/json\n'
        '\n'
        '{\n'
        '  "name": "Pushok",\n'
        '  "color": "white",\n'
        '  "birth_year": 2019,\n'
        '  "description": "Усатый, полосатый"\n'
        '}',
    )
    add_para(doc, 'Ответ — HTTP 201 Created:', indent=False)
    add_code(
        doc,
        '{\n'
        '  "id": 4,\n'
        '  "owner": "alice",\n'
        '  "name": "Pushok",\n'
        '  "color": "white",\n'
        '  "birth_year": 2019,\n'
        '  "description": "Усатый, полосатый",\n'
        '  "image": null,\n'
        '  "tags": [],\n'
        '  "achievements": [],\n'
        '  "created_at": "2026-05-03T15:24:02.539637+03:00"\n'
        '}',
    )

    add_heading_left(doc, 'Ошибка валидации (400)', level=3)
    add_code(
        doc,
        'POST /api/v1/cats/  body={"name":"BadYear","color":"black","birth_year":1800}\n'
        'HTTP 400 Bad Request\n'
        '{\n'
        '  "birth_year": ["Убедитесь, что это значение больше либо равно 1990."]\n'
        '}',
    )

    add_heading_left(doc, 'Запрет доступа к чужому ресурсу (403)', level=3)
    add_code(
        doc,
        'PATCH /api/v1/cats/1/  Authorization: Bearer <bob_access>\n'
        '(кот id=1 принадлежит alice)\n'
        'HTTP 403 Forbidden\n'
        '{ "detail": "У вас недостаточно прав для выполнения данного действия." }',
    )

    add_image(doc, SCREENS / '04_cats_list_json.png', width_cm=15,
              caption='Рисунок 2 — JSON-ответ GET /api/v1/cats/ (три кота, кириллица, теги, достижения)')

    # 4
    add_heading_left(doc, '4. Ключевые решения реализации (DRF)')
    add_para(
        doc,
        'CatViewSet наследует ModelViewSet и подключён через DefaultRouter в cats/urls.py. '
        'Для частных операций добавлен кастомный action upload_image с парсерами '
        'MultiPartParser и FormParser. AchievementViewSet и TagViewSet используют '
        'разные политики: первый — ModelViewSet с динамическим permission_classes '
        '(read public, write — IsAdminUser), второй — ReadOnlyModelViewSet.',
    )
    add_para(
        doc,
        'Сериализаторы написаны через ModelSerializer. CatSerializer возвращает '
        'вложенные TagSerializer и AchievementSerializer для чтения и принимает '
        'tag_ids/achievement_ids для записи (PrimaryKeyRelatedField). Поле owner '
        'выводится как username (SlugRelatedField). Кастомная валидация '
        'validate_name проверяет уникальность клички в рамках одного владельца '
        'на уровне сериализатора.',
    )
    add_para(
        doc,
        'Фильтрация настроена через django-filter (CatFilter с фильтрами owner, tag, '
        'birth_year_min, birth_year_max), поиск — через SearchFilter по полям '
        'name/color/description, сортировка — через OrderingFilter '
        '(created_at, birth_year, name). Пагинация — PageNumberPagination, '
        'page_size = 10, max_page_size = 100.',
    )
    add_para(
        doc,
        'JWT-аутентификация: настроен SIMPLE_JWT с lifetime access = 60 минут, '
        'refresh = 7 дней, тип заголовка — Bearer. URL /api/v1/auth/users/ и '
        '/api/v1/auth/users/me/ предоставлены Djoser, /api/v1/auth/jwt/{create,refresh,verify}/ — '
        'связкой Djoser + simplejwt. Throttling: anon 50/час, user 1000/час.',
    )

    # 5
    add_heading_left(doc, '5. Контроль доступа и обработка ошибок')
    add_para(
        doc,
        'Permissions: глобально IsAuthenticatedOrReadOnly (REST_FRAMEWORK '
        'настройка), на CatViewSet дополнительно — кастомный IsOwnerOrReadOnly '
        '(в core/permissions.py), который проверяет в has_object_permission, '
        'что obj.owner_id == request.user.id. На AchievementViewSet '
        'permission_classes выбираются динамически: GET — IsAuthenticatedOrReadOnly, '
        'остальные — IsAdminUser.',
    )
    add_table(
        doc,
        ['Сценарий', 'Запрос', 'Код', 'Ответ (фрагмент)'],
        [
            ['Анонимный POST на /cats/',
             'POST /api/v1/cats/ без Authorization',
             '401',
             '{"detail":"Учётные данные не были предоставлены."}'],
            ['Чужой кот PATCH',
             'PATCH /api/v1/cats/1/  bearer=bob (cat.owner=alice)',
             '403',
             '{"detail":"У вас недостаточно прав для выполнения данного действия."}'],
            ['Невалидный birth_year',
             'POST /api/v1/cats/  birth_year=1800',
             '400',
             '{"birth_year":["...больше либо равно 1990."]}'],
        ],
        col_widths_cm=[3.5, 5, 1.5, 6],
    )

    # 6
    add_heading_left(doc, '6. Документация и проверка запросами')
    add_para(
        doc,
        'API-документация генерируется drf-spectacular на основе вьюсетов и '
        'сериализаторов. Доступны три endpoint-а: '
        '/api/schema/ (OpenAPI 3.0 в формате YAML), '
        '/api/schema/swagger-ui/ (интерактивный Swagger UI с возможностью '
        '«try it out») и /api/schema/redoc/ (ReDoc — удобный для чтения '
        'трёхпанельный вид).',
    )
    add_image(doc, SCREENS / '02_swagger_authorize.png', width_cm=14,
              caption='Рисунок 3 — Swagger UI: диалог Authorize с jwtAuth (Bearer-схема, прокинута в OpenAPI 3.0)')
    add_image(doc, SCREENS / '03a_redoc_top.png', width_cm=14,
              caption='Рисунок 4 — ReDoc: заголовок Kittygram API + первый эндпоинт /achievements/ (часть 1 из 3)')
    add_image(doc, SCREENS / '03b_redoc_cats.png', width_cm=14,
              caption='Рисунок 5 — ReDoc: эндпоинты /cats/list и /cats/create — параметры запроса и схема тела (часть 2 из 3)')
    add_image(doc, SCREENS / '03c_redoc_schemas.png', width_cm=14,
              caption='Рисунок 6 — ReDoc: эндпоинт /tags/{id}/ + полный sidebar со всеми операциями API (часть 3 из 3)')

    add_heading_left(doc, '6.1. Постман-коллекция и newman')
    add_para(
        doc,
        'Коллекция postman/kittygram.postman_collection.json (формат v2.1) '
        'содержит 13 запросов: регистрация, JWT, профиль, list/create/get/patch/delete '
        'котов, upload_image, search+filter, негативные кейсы (401, 400). '
        'В коллекционные переменные записываются access_token и refresh_token '
        'из ответа JWT, далее они применяются в Bearer-авторизации. Прогон через '
        'newman показывает 13/13 успешных запросов и 6/6 пройденных тестов.',
    )
    add_code(
        doc,
        '$ newman run postman/kittygram.postman_collection.json \\\n'
        '       --env-var "base_url=http://127.0.0.1:8000"\n'
        '...\n'
        '│              iterations │   1 │   0 │\n'
        '│                requests │  13 │   0 │\n'
        '│            test-scripts │   6 │   0 │\n'
        '│              assertions │   6 │   0 │\n'
        '│ total run duration: 1.4s ...                                     │',
    )
    add_image(doc, SCREENS / '06_newman_report.png', width_cm=15,
              caption='Рисунок 7 — Newman HTML-репорт: все 13 запросов, 6/6 ассертов зелёные')

    add_heading_left(doc, '6.2. Примеры запросов через curl и httpie')
    add_code(
        doc,
        '# curl: получение JWT\n'
        'curl -X POST http://localhost:8000/api/v1/auth/jwt/create/ \\\n'
        '     -H "Content-Type: application/json" \\\n'
        '     -d \'{"username":"alice","password":"P@ssw0rd!2026Long"}\'\n'
        '\n'
        '# curl: список котов с фильтром и поиском\n'
        'curl "http://localhost:8000/api/v1/cats/?owner=alice&search=пуш&ordering=-created_at"\n'
        '\n'
        '# httpie: создание кота с авторизацией\n'
        'http POST :8000/api/v1/cats/ \\\n'
        '     "Authorization:Bearer $TOKEN" \\\n'
        '     name=Murka color=серый birth_year:=2022',
    )

    # 7
    add_heading_left(doc, '7. Деплой и Docker')
    add_para(
        doc,
        'Корневой Dockerfile собирает образ на python:3.12-slim: устанавливает '
        'libpq-dev для psycopg2, ставит зависимости, копирует исходный код, '
        'выполняет collectstatic при старте контейнера. Запуск через gunicorn '
        'kittygram.wsgi:application на 0.0.0.0:8000 с 3 worker-ами.',
    )
    add_para(
        doc,
        'docker-compose.yml объединяет три сервиса: db (postgres:16-alpine с '
        'pg_data volume и healthcheck), backend (build из Dockerfile, '
        'depends_on db по condition: service_healthy, монтирует static_volume '
        'и media_volume) и nginx (через профиль with-nginx — reverse-proxy на '
        'backend, отдаёт статику и медиа). Все секреты вынесены в .env, '
        'образец — .env.example в репозитории.',
    )
    add_code(
        doc,
        '# Запуск (быстрый старт)\n'
        'cp .env.example .env       # затем заполнить SECRET_KEY и POSTGRES_PASSWORD\n'
        'docker compose up -d --build\n'
        'docker compose exec backend python manage.py createsuperuser\n'
        '\n'
        '# Проверка\n'
        'curl http://localhost:8000/api/v1/cats/        # → {"count":0,"next":null,...}\n'
        '\n'
        '# Остановка\n'
        'docker compose down            # сохранить volumes\n'
        'docker compose down -v         # снести вместе с БД и медиа',
    )
    add_image(doc, SCREENS / '05_admin_cats_list.png', width_cm=15,
              caption='Рисунок 8 — Django-админка под docker-compose: список котов (Пушок, Барсик, Мурзик) из PostgreSQL, видны фильтры и иерархия моделей')

    # 8
    add_heading_left(doc, '8. Итоги и самооценка')
    add_para(doc,
             '1. API-стек обновлён до текущих версий (Django 4.2 LTS, DRF 3.14), '
             'переведён с SQLite на PostgreSQL, разворачивается одной командой '
             'docker compose up -d --build.',
             indent=True)
    add_para(doc,
             '2. Реализована полная схема контроля доступа: глобальные DRF '
             'permissions, кастомный IsOwnerOrReadOnly, JWT через Djoser+simplejwt, '
             'throttling. Доказано тестами: 401 на анонимный POST, 403 на чужой '
             'ресурс, 400 на нарушения валидаций.',
             indent=True)
    add_para(doc,
             '3. Сложнее всего было корректно собрать сериализаторы со связями '
             'M2M через through-модели (CatTag/CatAchievement) так, чтобы на '
             'чтение возвращались вложенные объекты, а на запись принимались '
             'списки идентификаторов.',
             indent=True)
    add_para(doc,
             '4. В следующей версии стоит добавить покрытие автотестами '
             '(pytest-django + APIClient), CI на GitHub Actions, а также '
             'настроить ротацию refresh-токенов и blacklist для logout.',
             indent=True)
    add_para(doc,
             '5. Архитектурно проект подготовлен к расширению: уже выделены '
             'core/permissions.py и core/pagination.py, что позволит '
             'переиспользовать их в творческом задании (мини-чат по заявкам '
             'на кото-ивенты).',
             indent=True)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    sys.stdout.reconfigure(encoding='utf-8')
    print(f'OK: {OUT} ({OUT.stat().st_size} bytes)')


if __name__ == '__main__':
    build()
