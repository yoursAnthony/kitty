"""Точечная замена шапки в существующем docx на полноценный титульный лист по РЭУ.

Удаляет «черновую шапку» (всё до первого нумерованного раздела вида «N.» / «N.N»)
и вставляет в начало:
    - блок Минобрнауки + РЭУ им. Плеханова (центр, заглавные)
    - институт / кафедра (заглушки или переданные значения)
    - тип работы (творческое задание / отчёт / курсовая)
    - дисциплина и тема
    - блок «Выполнил / Проверил»
    - «Москва, год»
    - разрыв страницы

Запуск (из корня проекта):
    .venv/Scripts/python scripts/insert_title_page.py
"""
import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def _set_run_font(run, size=14, bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')


def _para(doc, text, bold=False, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
          line_spacing=1.0, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        _set_run_font(run, size=size, bold=bold)
    return p


def build_title_page(work_kind: str, discipline: str, topic: str,
                     institute: str, department: str,
                     student: str, group: str, supervisor: str, year: int):
    """Собирает Document только с титульником + page break и возвращает его.

    Спроектирован, чтобы помещаться в одну страницу A4 при полях 30/15/20/20.
    """
    doc = Document()

    # Поля как в основном документе (РЭУ)
    section = doc.sections[0]
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Шапка: Министерство + ВУЗ — компактно (одинарный интервал)
    _para(doc, 'МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ',
          bold=True, size=12)
    _para(doc, 'федеральное государственное бюджетное образовательное учреждение '
               'высшего образования', size=12, space_before=4)
    _para(doc, '«РОССИЙСКИЙ ЭКОНОМИЧЕСКИЙ УНИВЕРСИТЕТ имени Г. В. ПЛЕХАНОВА»',
          bold=True, size=12)

    # Институт / кафедра — компактный отступ
    _para(doc, institute, size=12, space_before=14)
    _para(doc, department, size=12)

    # Тип работы — большой отступ от шапки + крупный шрифт
    _para(doc, work_kind.upper(), bold=True, size=22, space_before=72)

    # Дисциплина и тема — без лишних пустых параграфов
    _para(doc, 'по дисциплине', size=14, space_before=18)
    _para(doc, f'«{discipline}»', bold=True, size=14)
    _para(doc, 'на тему:', size=14, space_before=10)
    _para(doc, f'«{topic}»', bold=True, size=14)

    # Блок «Выполнил/Проверил» — справа, компактный
    _para(doc, 'Выполнил:', bold=True, size=14,
          align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=72)
    _para(doc, f'студент группы {group}', size=14, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(doc, student, size=14, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(doc, 'Проверил:', bold=True, size=14,
          align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=10)
    _para(doc, supervisor, size=14, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Низ страницы — Москва, год
    _para(doc, f'Москва, {year}', size=14, space_before=42)

    # Разрыв страницы
    page_break_p = doc.add_paragraph()
    page_break_p.paragraph_format.line_spacing = 1.0
    page_break_p.paragraph_format.space_before = Pt(0)
    page_break_p.paragraph_format.space_after = Pt(0)
    run = page_break_p.add_run()
    run.add_break(WD_BREAK.PAGE)

    return doc


def replace_title_page(target_path: Path, work_kind, discipline, topic,
                       institute, department, student, group, supervisor, year,
                       cut_marker_predicate):
    """Удаляет «черновую шапку» из target docx и вставляет полноценный титульник.

    cut_marker_predicate(text) -> bool: первая параграф, для которого True, считается
    началом основного содержимого; всё ВЫШЕ него удаляется.
    """
    doc = Document(str(target_path))
    body = doc.element.body

    paragraphs = doc.paragraphs
    cut_index = None
    for i, p in enumerate(paragraphs):
        if cut_marker_predicate(p.text.strip()):
            cut_index = i
            break
    if cut_index is None:
        raise RuntimeError(
            f'Не нашёл маркер начала контента в {target_path.name} — '
            'проверьте предикат.'
        )

    # Удаляем «черновую шапку»
    for p in paragraphs[:cut_index]:
        el = p._element
        el.getparent().remove(el)

    # Собираем титульник в helper-документе
    title_doc = build_title_page(
        work_kind=work_kind, discipline=discipline, topic=topic,
        institute=institute, department=department,
        student=student, group=group, supervisor=supervisor, year=year,
    )

    # Берём все block-элементы из helper (paragraphs), кроме sectPr
    title_blocks = []
    for el in title_doc.element.body:
        if el.tag == qn('w:sectPr'):
            continue
        title_blocks.append(el)

    # Вставляем их В НАЧАЛО target body перед первым оставшимся элементом
    first_existing = body[0]
    for el in title_blocks:
        first_existing.addprevious(el)

    doc.save(str(target_path))
    sys.stdout.reconfigure(encoding='utf-8')
    return cut_index


# ─────── Конфигурации для каждого документа ───────

REPO = Path(__file__).resolve().parent.parent
OUT_DIR = (
    REPO.parent.parent.parent
    / 'Any' / 'PREU'
    / 'интеграция и управление приложениями на удаленном сервере'
)

DEFAULTS = {
    'institute': '[ИНСТИТУТ / ФАКУЛЬТЕТ ___________________________]',
    'department': '[Кафедра _____________________________________]',
    'discipline': 'Интеграция и управление приложениями на удалённом сервере',
    'student': 'Сергеев А. А.',
    'group': 'ПИ2у/24б',
    'supervisor': 'Брызгалов А. А.',
    'year': 2026,
}

DOCS = {
    'creative': {
        'path': OUT_DIR / 'Творческое_задание.docx',
        'work_kind': 'Творческое задание',
        'topic': 'Проектирование и реализация серверной части проекта Kittygram '
                 'для поддержки пользовательского сценария «Мини-чат по заявкам»',
        'cut_marker': lambda t: t.startswith('1. Выбор темы'),
    },
    'current': {
        'path': OUT_DIR / 'Отчёт_текущий_рейтинг.docx',
        'work_kind': 'Отчёт по выполнению практикума',
        'topic': 'Разработка REST API на Django REST Framework. Проект Kittygram',
        'cut_marker': lambda t: t.startswith('1. Цель и итоговый результат'),
    },
    'kursovaya': {
        'path': OUT_DIR / 'Курсовая_работа.docx',
        'work_kind': 'Курсовая работа',
        'topic': 'Проектирование и реализация серверной части проекта Kittygram '
                 'для поддержки пользовательского сценария «Мини-чат по заявкам»',
        'cut_marker': lambda t: t.startswith('Введение') or t.startswith('1.'),
    },
}


def apply_to(doc_key: str, **overrides):
    cfg = DOCS[doc_key]
    if not cfg['path'].exists():
        print(f'[skip] {cfg["path"].name} — файла нет')
        return
    params = {**DEFAULTS, **overrides}
    cut = replace_title_page(
        target_path=cfg['path'],
        work_kind=cfg['work_kind'],
        topic=cfg['topic'],
        cut_marker_predicate=cfg['cut_marker'],
        **params,
    )
    sys.stdout.reconfigure(encoding='utf-8')
    print(f'[ok] {cfg["path"].name}: вырезано {cut} старых параграфов, '
          'добавлен титульник.')


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument(
        'doc', choices=list(DOCS.keys()) + ['all'],
        help='какой документ переоформить',
    )
    parser.add_argument('--institute', default=DEFAULTS['institute'])
    parser.add_argument('--department', default=DEFAULTS['department'])
    args = parser.parse_args()

    overrides = {'institute': args.institute, 'department': args.department}
    targets = list(DOCS.keys()) if args.doc == 'all' else [args.doc]
    for k in targets:
        apply_to(k, **overrides)
